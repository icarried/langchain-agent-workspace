from __future__ import annotations

import re
from pathlib import Path

from docx import Document
from docx.shared import Mm

from .schemas import DocumentElement, FormatFinding


def inspect_official_document(path: str | Path, elements: list[DocumentElement]) -> list[FormatFinding]:
    findings: list[FormatFinding] = []
    document_path = Path(path)
    if not elements:
        findings.append(
            FormatFinding(
                check_id="content-empty",
                severity="high",
                category="正文内容",
                message="未解析到可检查的公文正文。",
                suggestion="请确认文件不是扫描件、图片或空文档。",
            )
        )
        return findings

    findings.extend(_inspect_text_structure(elements))
    if document_path.suffix.lower() == ".docx":
        findings.extend(_inspect_docx_page_setup(document_path))
        findings.extend(_inspect_docx_runs(document_path))
    else:
        findings.append(
            FormatFinding(
                check_id="layout-source-limited",
                severity="info",
                category="版式检查",
                message="当前输入不是 DOCX，无法读取页边距、纸张和字体等版式元数据。",
                suggestion="如需完整核验 GB/T 9704-2012 版式，请提供原始 DOCX 文件。",
            )
        )
    return findings


def _inspect_text_structure(elements: list[DocumentElement]) -> list[FormatFinding]:
    findings: list[FormatFinding] = []
    texts = [element.text for element in elements]
    first = texts[0]
    joined = "\n".join(texts)

    if len(first) > 80:
        findings.append(
            FormatFinding(
                check_id="title-too-long",
                severity="medium",
                category="标题",
                message="首段标题疑似过长，可能混入正文或说明文字。",
                suggestion="将标题控制为准确、简洁的公文标题，正文另起段落。",
                evidence=f"paragraph#1: {first[:80]}",
            )
        )

    if not _has_issuer_or_recipient(texts):
        findings.append(
            FormatFinding(
                check_id="missing-recipient",
                severity="medium",
                category="主送机关",
                message="未识别到明确的主送机关或收文对象。",
                suggestion="请在标题下方补充主送机关，并使用全称或规范简称。",
            )
        )

    if not re.search(r"\d{4}\s*年\s*\d{1,2}\s*月\s*\d{1,2}\s*日", joined):
        findings.append(
            FormatFinding(
                check_id="missing-date",
                severity="medium",
                category="成文日期",
                message="未识别到规范的中文成文日期。",
                suggestion="在落款处使用“YYYY年M月D日”格式标注成文日期。",
            )
        )

    if not re.search(r"(特此|妥否|请示|报告|通知|决定|意见|函)", joined):
        findings.append(
            FormatFinding(
                check_id="genre-marker-weak",
                severity="low",
                category="文种与结束语",
                message="未识别到明显文种或常用公文结束语，文种边界可能不清晰。",
                suggestion="根据实际用途明确文种，并补充与文种匹配的结束语或办理要求。",
            )
        )

    if not findings:
        findings.append(
            FormatFinding(
                check_id="text-structure-pass",
                severity="info",
                category="正文结构",
                message="文本结构初检未发现明显缺漏。",
                suggestion="正式定稿前仍应结合发文机关模板复核版记、附件和印章页。",
            )
        )
    return findings


def _has_issuer_or_recipient(texts: list[str]) -> bool:
    candidates = texts[:5]
    return any(text.endswith(("：", ":")) or "公司" in text or "局" in text or "委" in text for text in candidates[1:])


def _inspect_docx_page_setup(path: Path) -> list[FormatFinding]:
    document = Document(path)
    findings: list[FormatFinding] = []
    if not document.sections:
        return findings
    section = document.sections[0]
    expected = {
        "top_margin": Mm(37),
        "bottom_margin": Mm(35),
        "left_margin": Mm(28),
        "right_margin": Mm(26),
    }
    actual = {
        "top_margin": section.top_margin,
        "bottom_margin": section.bottom_margin,
        "left_margin": section.left_margin,
        "right_margin": section.right_margin,
    }
    labels = {
        "top_margin": "上边距",
        "bottom_margin": "下边距",
        "left_margin": "左边距",
        "right_margin": "右边距",
    }
    for key, expected_value in expected.items():
        value = actual[key]
        if value is None:
            continue
        if abs(value.mm - expected_value.mm) > 2:
            findings.append(
                FormatFinding(
                    check_id=f"page-{key}",
                    severity="medium",
                    category="页面版式",
                    message=f"{labels[key]}约为 {value.mm:.1f} mm，与 GB/T 9704-2012 常用设置不一致。",
                    suggestion=f"建议将{labels[key]}调整为约 {expected_value.mm:.0f} mm，并结合单位模板复核。",
                )
            )

    if section.page_width and section.page_height:
        width = section.page_width.mm
        height = section.page_height.mm
        if abs(width - 210) > 3 or abs(height - 297) > 3:
            findings.append(
                FormatFinding(
                    check_id="page-size",
                    severity="high",
                    category="页面版式",
                    message=f"纸张尺寸约为 {width:.1f} mm x {height:.1f} mm，疑似不是 A4。",
                    suggestion="公文通常应使用 A4 纸张，建议调整页面纸张大小。",
                )
            )
    return findings


def _inspect_docx_runs(path: Path) -> list[FormatFinding]:
    document = Document(path)
    findings: list[FormatFinding] = []
    non_empty = [paragraph for paragraph in document.paragraphs if paragraph.text.strip()]
    if not non_empty:
        return findings

    first = non_empty[0]
    title_font_names = {run.font.name for run in first.runs if run.text.strip() and run.font.name}
    if title_font_names and not any("方正小标宋" in name or "宋" in name for name in title_font_names):
        findings.append(
            FormatFinding(
                check_id="title-font",
                severity="low",
                category="字体",
                message=f"标题显式字体为 {', '.join(sorted(title_font_names))}，可能不符合常用公文标题字体。",
                suggestion="标题通常采用方正小标宋简体或单位模板指定字体，二号字居中。",
                evidence="paragraph#1",
            )
        )

    body_font_names = set()
    for paragraph in non_empty[1:8]:
        for run in paragraph.runs:
            if run.text.strip() and run.font.name:
                body_font_names.add(run.font.name)
    if body_font_names and not any("仿宋" in name for name in body_font_names):
        findings.append(
            FormatFinding(
                check_id="body-font",
                severity="low",
                category="字体",
                message=f"正文显式字体为 {', '.join(sorted(body_font_names))}，未识别到仿宋类字体。",
                suggestion="正文通常采用仿宋_GB2312 或单位模板指定仿宋类字体，三号字。",
            )
        )
    return findings

