from __future__ import annotations

import os
import re
import socket
import xml.etree.ElementTree as ET
import zipfile
from io import BytesIO
from pathlib import Path
from urllib.error import HTTPError, URLError
from urllib.parse import unquote, urlsplit
from urllib.request import Request, urlopen

from docx import Document
from pypdf import PdfReader

from .doc_converter import convert_doc_to_docx
from .ocr import ocr_image_bytes
from .schemas import ResumeElement


SUPPORTED_EXTENSIONS = {".doc", ".docx", ".md", ".pdf", ".txt"}
DEFAULT_OCR_MAX_PAGES = 50
DEFAULT_MIN_REVIEWABLE_TEXT_CHARS = 12
OCR_IMAGE_MIME_TYPES = {
    ".bmp": "image/bmp",
    ".jpeg": "image/jpeg",
    ".jpg": "image/jpeg",
    ".png": "image/png",
    ".tif": "image/tiff",
    ".tiff": "image/tiff",
    ".webp": "image/webp",
}
DEFAULT_REMOTE_FILE_LIMIT_BYTES = 10 * 1024 * 1024
DEFAULT_REMOTE_TIMEOUT_SECONDS = 30.0
REMOTE_CHUNK_SIZE = 64 * 1024


def load_resume_elements(path: str | Path) -> list[ResumeElement]:
    if is_remote_resume_source(path):
        return _load_remote_resume(str(path))

    resume_path = Path(path)
    if not resume_path.exists():
        raise FileNotFoundError(f"resume file not found: {resume_path}")

    suffix = resume_path.suffix.lower()
    if suffix in {".md", ".txt"}:
        return _load_txt(resume_path)
    if suffix == ".doc":
        return _load_doc(resume_path)
    if suffix == ".docx":
        return _load_docx(resume_path)
    if suffix == ".pdf":
        return _load_pdf(resume_path)

    supported = ", ".join(sorted(SUPPORTED_EXTENSIONS))
    raise ValueError(f"unsupported resume file type '{suffix}', supported: {supported}")


def is_remote_resume_source(source: str | Path) -> bool:
    if isinstance(source, Path):
        return False
    return urlsplit(str(source)).scheme.lower() in {"http", "https"}


def resume_source_filename(source: str | Path) -> str:
    if not is_remote_resume_source(source):
        return Path(source).name
    parsed = urlsplit(str(source))
    filename = Path(unquote(parsed.path)).name
    if not filename:
        raise ValueError("remote resume URL path must contain a filename")
    return filename


def safe_resume_source_label(source: str | Path) -> str:
    if not is_remote_resume_source(source):
        return str(source)
    parsed = urlsplit(str(source))
    return parsed._replace(query="", fragment="").geturl()


def _load_remote_resume(url: str) -> list[ResumeElement]:
    filename = resume_source_filename(url)
    suffix = Path(filename).suffix.lower()
    _validate_supported_suffix(suffix)
    _validate_remote_url(url)
    data = _download_remote_bytes(url)
    stream = BytesIO(data)
    if suffix in {".md", ".txt"}:
        try:
            text = data.decode("utf-8-sig")
        except UnicodeDecodeError as exc:
            raise ValueError(
                f"remote resume '{filename}' is not valid UTF-8 text"
            ) from exc
        return _elements_from_lines(
            text.splitlines(), kind="paragraph", source=filename
        )
    if suffix == ".doc":
        return _load_doc(stream, source=filename)
    if suffix == ".docx":
        return _load_docx(stream, source=filename)
    return _load_pdf(stream, source=filename)


def _download_remote_bytes(url: str) -> bytes:
    filename = resume_source_filename(url)
    limit = _positive_int_env(
        "BATCH_RESUME_REVIEW_MAX_REMOTE_FILE_BYTES", DEFAULT_REMOTE_FILE_LIMIT_BYTES
    )
    timeout = _positive_float_env(
        "BATCH_RESUME_REVIEW_REMOTE_TIMEOUT_SECONDS", DEFAULT_REMOTE_TIMEOUT_SECONDS
    )
    request = Request(url, headers={"User-Agent": "batch-resume-review-llm/0.1"})
    try:
        response, used_localhost_fallback = _open_remote(request, timeout)
        with response:
            if not used_localhost_fallback:
                _validate_remote_url(response.geturl())
            content_length = response.headers.get("Content-Length")
            if content_length and int(content_length) > limit:
                raise ValueError(
                    f"remote resume '{filename}' exceeds the {limit}-byte size limit"
                )
            chunks: list[bytes] = []
            received = 0
            while True:
                chunk = response.read(min(REMOTE_CHUNK_SIZE, limit - received + 1))
                if not chunk:
                    break
                received += len(chunk)
                if received > limit:
                    raise ValueError(
                        f"remote resume '{filename}' exceeds the {limit}-byte size limit"
                    )
                chunks.append(chunk)
    except HTTPError as exc:
        detail = _remote_http_error_detail(exc)
        raise ValueError(
            f"failed to download remote resume '{filename}': HTTP {exc.code}{detail}"
        ) from exc
    except (URLError, TimeoutError, socket.timeout) as exc:
        raise ValueError(
            f"failed to download remote resume '{filename}': network error"
        ) from exc
    return b"".join(chunks)


def _open_remote(request: Request, timeout: float) -> tuple[object, bool]:
    try:
        response = urlopen(request, timeout=timeout)  # noqa: S310 - URL is validated.
        return response, False
    except HTTPError:
        raise
    except (URLError, TimeoutError, socket.timeout) as original_error:
        fallback_request = _localhost_fallback_request(request.full_url)
        if fallback_request is None:
            raise original_error
        response = urlopen(  # noqa: S310 - fallback is restricted to localhost.
            fallback_request,
            timeout=timeout,
        )
        return response, True


def _localhost_fallback_request(url: str) -> Request | None:
    """Build the FastGPT/WSL compatibility request without changing signed Host.

    FastGPT may sign a Windows LAN origin such as 10.71.2.94:9000 while its
    MinIO container is only reachable from Windows through a localhost-published
    port (currently 127.0.0.1:9002). AWS V4 includes Host in the signature, so
    rewriting the public URL itself would invalidate the signature. This helper
    changes only the TCP transport origin and deliberately preserves Host.

    This is a deployment compatibility layer, not general retry behavior. The
    preferred long-term fix is to make FastGPT sign its actual externally
    reachable MinIO origin; see this agent's README deployment notes.
    """
    parsed = urlsplit(url)
    if parsed.scheme.lower() != "http" or not parsed.hostname:
        return None
    try:
        local_addresses = set(socket.gethostbyname_ex(socket.gethostname())[2])
    except OSError:
        return None
    if parsed.hostname not in local_addresses:
        return None
    transport_netloc = _local_minio_transport_netloc(parsed.port or 80)
    transport_url = parsed._replace(netloc=transport_netloc).geturl()
    return Request(
        transport_url,
        headers={
            "User-Agent": "batch-resume-review-llm/0.1",
            "Host": parsed.netloc,
        },
    )


def _local_minio_transport_netloc(original_port: int) -> str:
    configured = os.getenv("BATCH_RESUME_REVIEW_LOCAL_MINIO_ENDPOINT", "").strip()
    if not configured:
        return f"127.0.0.1:{original_port}"
    parsed = urlsplit(configured)
    if (
        parsed.scheme.lower() != "http"
        or parsed.hostname not in {"127.0.0.1", "localhost"}
        or parsed.username
        or parsed.password
        or parsed.path not in {"", "/"}
        or parsed.query
        or parsed.fragment
    ):
        raise ValueError(
            "BATCH_RESUME_REVIEW_LOCAL_MINIO_ENDPOINT must be an HTTP localhost origin"
        )
    return parsed.netloc


def _validate_remote_url(url: str) -> None:
    parsed = urlsplit(url)
    if parsed.scheme.lower() not in {"http", "https"} or not parsed.hostname:
        raise ValueError("remote resume source must be an HTTP(S) URL")
    if parsed.username or parsed.password:
        raise ValueError("remote resume URL must not contain user information")
    allowed = {
        item.strip().lower()
        for item in os.getenv("BATCH_RESUME_REVIEW_ALLOWED_URL_HOSTS", "").split(",")
        if item.strip()
    }
    if allowed and parsed.hostname.lower() not in allowed:
        raise ValueError(f"remote resume host '{parsed.hostname}' is not allowed")


def _remote_http_error_detail(error: HTTPError) -> str:
    code = ""
    try:
        root = ET.fromstring(error.read(8192))
        candidate = (root.findtext("Code") or "").strip()
        if re.fullmatch(r"[A-Za-z][A-Za-z0-9]{0,63}", candidate):
            code = candidate
    except (ET.ParseError, OSError, ValueError):
        pass

    request_id = ""
    if error.headers:
        candidate = (error.headers.get("X-Amz-Request-Id") or "").strip()
        if re.fullmatch(r"[A-Za-z0-9-]{1,128}", candidate):
            request_id = candidate

    details = [
        value
        for value in (code, f"request_id={request_id}" if request_id else "")
        if value
    ]
    return f" ({', '.join(details)})" if details else ""


def _validate_supported_suffix(suffix: str) -> None:
    if suffix not in SUPPORTED_EXTENSIONS:
        supported = ", ".join(sorted(SUPPORTED_EXTENSIONS))
        raise ValueError(
            f"unsupported resume file type '{suffix}', supported: {supported}"
        )


def _positive_int_env(name: str, default: int) -> int:
    try:
        value = int(os.getenv(name, str(default)))
    except ValueError as exc:
        raise ValueError(f"{name} must be a positive integer") from exc
    if value <= 0:
        raise ValueError(f"{name} must be a positive integer")
    return value


def _positive_float_env(name: str, default: float) -> float:
    try:
        value = float(os.getenv(name, str(default)))
    except ValueError as exc:
        raise ValueError(f"{name} must be a positive number") from exc
    if value <= 0:
        raise ValueError(f"{name} must be a positive number")
    return value


def _load_txt(path: Path) -> list[ResumeElement]:
    text = path.read_text(encoding="utf-8-sig")
    return _elements_from_lines(text.splitlines(), kind="paragraph", source=path.name)


def _load_doc(
    path: Path | BytesIO, *, source: str | None = None
) -> list[ResumeElement]:
    data = _source_bytes(path)
    source_name = source or Path(path).name
    converted = convert_doc_to_docx(data, source=source_name)
    return _load_docx(BytesIO(converted), source=source_name)


def _load_docx(
    path: Path | BytesIO, *, source: str | None = None
) -> list[ResumeElement]:
    document = Document(path)
    source_name = source or Path(path).name
    elements: list[ResumeElement] = []
    for paragraph in document.paragraphs:
        text = paragraph.text.strip()
        if not text:
            continue
        elements.append(
            ResumeElement(
                index=len(elements) + 1,
                kind="paragraph",
                text=text,
                source=source_name,
                style=paragraph.style.name if paragraph.style else "",
            )
        )

    for table in document.tables:
        rows = []
        for row in table.rows:
            cells = [cell.text.strip().replace("\n", " ") for cell in row.cells]
            if any(cells):
                rows.append(" | ".join(cells))
        if rows:
            elements.append(
                ResumeElement(
                    index=len(elements) + 1,
                    kind="table",
                    text="\n".join(rows),
                    source=source_name,
                )
            )

    if _has_reviewable_elements(elements):
        return elements

    images = _docx_images(_source_bytes(path))
    if not images:
        raise ValueError(
            f"DOCX has insufficient extractable text and no OCR-compatible images: {source_name}"
        )
    return elements + _ocr_images(
        images, source=source_name, start_index=len(elements) + 1
    )


def _load_pdf(
    path: Path | BytesIO, *, source: str | None = None
) -> list[ResumeElement]:
    data = _source_bytes(path)
    reader = PdfReader(BytesIO(data))
    source_name = source or Path(path).name
    elements: list[ResumeElement] = []
    ocr_pages = 0
    max_ocr_pages = _positive_int_env(
        "BATCH_RESUME_REVIEW_OCR_MAX_PAGES", DEFAULT_OCR_MAX_PAGES
    )
    for page_number, page in enumerate(reader.pages, start=1):
        page_text = page.extract_text() or ""
        if _has_reviewable_text(page_text):
            for line in _clean_lines(page_text.splitlines()):
                elements.append(
                    ResumeElement(
                        index=len(elements) + 1,
                        kind="pdf_line",
                        text=line,
                        source=f"{source_name}:page-{page_number}",
                    )
                )
            continue

        ocr_pages += 1
        if ocr_pages > max_ocr_pages:
            raise ValueError(
                f"PDF requires OCR on more than {max_ocr_pages} pages: {source_name}"
            )
        image = _render_pdf_page(data, page_number - 1, source=source_name)
        text = ocr_image_bytes(
            image,
            "image/png",
            source=f"{source_name}:page-{page_number}",
        )
        elements.extend(
            _elements_from_lines(
                text.splitlines(),
                kind="ocr_line",
                source=f"{source_name}:page-{page_number}",
                start_index=len(elements) + 1,
            )
        )
    if not elements:
        raise ValueError(f"PDF contains no reviewable text: {source_name}")
    return elements


def _source_bytes(path: Path | BytesIO) -> bytes:
    if isinstance(path, BytesIO):
        return path.getvalue()
    return path.read_bytes()


def _docx_images(data: bytes) -> list[tuple[bytes, str, str]]:
    images = []
    try:
        with zipfile.ZipFile(BytesIO(data)) as archive:
            for name in archive.namelist():
                suffix = Path(name).suffix.lower()
                mime_type = OCR_IMAGE_MIME_TYPES.get(suffix)
                if name.startswith("word/media/") and mime_type:
                    images.append((archive.read(name), mime_type, Path(name).name))
    except zipfile.BadZipFile as exc:
        raise ValueError("DOCX file is corrupt or invalid") from exc
    return images


def _ocr_images(
    images: list[tuple[bytes, str, str]],
    *,
    source: str,
    start_index: int,
) -> list[ResumeElement]:
    max_images = _positive_int_env(
        "BATCH_RESUME_REVIEW_OCR_MAX_PAGES", DEFAULT_OCR_MAX_PAGES
    )
    if len(images) > max_images:
        raise ValueError(
            f"document requires OCR on more than {max_images} images: {source}"
        )
    elements: list[ResumeElement] = []
    for data, mime_type, image_name in images:
        image_source = f"{source}:{image_name}"
        text = ocr_image_bytes(data, mime_type, source=image_source)
        elements.extend(
            _elements_from_lines(
                text.splitlines(),
                kind="ocr_line",
                source=image_source,
                start_index=start_index + len(elements),
            )
        )
    if not elements:
        raise ValueError(f"Bailian OCR returned no reviewable text for '{source}'")
    return elements


def _render_pdf_page(data: bytes, page_index: int, *, source: str) -> bytes:
    try:
        import fitz
    except ImportError as exc:
        raise ValueError(
            f"PDF OCR requires PyMuPDF, but it is not installed: {source}"
        ) from exc
    try:
        with fitz.open(stream=data, filetype="pdf") as document:
            page = document.load_page(page_index)
            pixmap = page.get_pixmap(matrix=fitz.Matrix(2, 2), alpha=False)
            return pixmap.tobytes("png")
    except Exception as exc:
        raise ValueError(f"failed to render PDF page for OCR: {source}") from exc


def _has_reviewable_elements(elements: list[ResumeElement]) -> bool:
    return _has_reviewable_text("\n".join(element.text for element in elements))


def _has_reviewable_text(text: str) -> bool:
    minimum = _positive_int_env(
        "BATCH_RESUME_REVIEW_MIN_TEXT_CHARS", DEFAULT_MIN_REVIEWABLE_TEXT_CHARS
    )
    count = sum(
        character.isalnum() or "\u4e00" <= character <= "\u9fff" for character in text
    )
    return count >= minimum


def _elements_from_lines(
    lines: list[str],
    *,
    kind: str,
    source: str,
    start_index: int = 1,
) -> list[ResumeElement]:
    return [
        ResumeElement(index=index, kind=kind, text=line, source=source)
        for index, line in enumerate(_clean_lines(lines), start=start_index)
    ]


def _clean_lines(lines: list[str]) -> list[str]:
    return [line.strip() for line in lines if line.strip()]
