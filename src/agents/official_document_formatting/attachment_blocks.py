"""Controlled normalization for attachment labels and item paragraphs."""

from __future__ import annotations

import re

from docx.oxml import OxmlElement
from docx.text.paragraph import Paragraph

from .roles import ParagraphRole, classify_paragraphs

_ATTACHMENT_LABEL_PATTERN = re.compile(r"^(附件[：:])\s*(.*)$", re.DOTALL)
_LEADING_LAYOUT_WHITESPACE = " \t\u3000"


def normalize_attachment_blocks(document) -> None:
    """Put the attachment label and list on separate, clean paragraphs."""
    paragraphs = list(document.paragraphs)
    roles = classify_paragraphs([paragraph.text for paragraph in paragraphs])
    for paragraph, role in zip(paragraphs, roles, strict=True):
        if role is ParagraphRole.ATTACHMENT_NOTE:
            match = _ATTACHMENT_LABEL_PATTERN.fullmatch(paragraph.text.strip())
            if match is None:
                continue
            label, first_item = match.groups()
            paragraph.text = label
            if first_item:
                _insert_paragraph_after(paragraph, first_item.lstrip(_LEADING_LAYOUT_WHITESPACE))
        elif role is ParagraphRole.ATTACHMENT_ITEM:
            paragraph.text = paragraph.text.lstrip(_LEADING_LAYOUT_WHITESPACE)


def semantic_paragraph_texts(document) -> tuple[str, ...]:
    """Normalize only the authorized attachment layout differences for comparison."""
    texts = [paragraph.text for paragraph in document.paragraphs]
    roles = classify_paragraphs(texts)
    result: list[str] = []
    for text, role in zip(texts, roles, strict=True):
        if role is ParagraphRole.ATTACHMENT_NOTE:
            match = _ATTACHMENT_LABEL_PATTERN.fullmatch(text.strip())
            if match is None:
                result.append(text)
                continue
            label, first_item = match.groups()
            result.append(label)
            if first_item:
                result.append(first_item.lstrip(_LEADING_LAYOUT_WHITESPACE))
        elif role is ParagraphRole.ATTACHMENT_ITEM:
            result.append(text.lstrip(_LEADING_LAYOUT_WHITESPACE))
        else:
            result.append(text)
    return tuple(result)


def _insert_paragraph_after(paragraph, text: str) -> Paragraph:
    paragraph_element = OxmlElement("w:p")
    paragraph._p.addnext(paragraph_element)
    inserted = Paragraph(paragraph_element, paragraph._parent)
    inserted.add_run(text)
    return inserted
