"""Page geometry, document grid, and odd/even page-number formatting."""

from __future__ import annotations

from docx.enum.section import WD_ORIENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement, parse_xml
from docx.oxml.ns import nsdecls, qn
from docx.shared import Cm

from .standards import (
    BOTTOM_MARGIN_CM,
    EXACT_LINE_SPACING_TWIPS,
    FOOTER_DISTANCE_CM,
    LATIN_FONT,
    LEFT_MARGIN_CM,
    ONE_CHAR_INDENT_CHARS,
    ONE_CHAR_INDENT_TWIPS,
    PAGE_HEIGHT_CM,
    PAGE_NUMBER_SIZE_PT,
    PAGE_WIDTH_CM,
    RIGHT_MARGIN_CM,
    SONG,
    TOP_MARGIN_CM,
)


def apply_page_layout(document) -> None:
    """Apply the normative A4 page geometry, grid, and dynamic page numbers."""
    _enable_odd_even_footers(document)
    _enable_field_updates(document)
    for section_index, section in enumerate(document.sections):
        section.orientation = WD_ORIENT.PORTRAIT
        section.page_width = Cm(PAGE_WIDTH_CM)
        section.page_height = Cm(PAGE_HEIGHT_CM)
        section.top_margin = Cm(TOP_MARGIN_CM)
        section.bottom_margin = Cm(BOTTOM_MARGIN_CM)
        section.left_margin = Cm(LEFT_MARGIN_CM)
        section.right_margin = Cm(RIGHT_MARGIN_CM)
        section.footer_distance = Cm(FOOTER_DISTANCE_CM)
        _set_document_grid(section)
        if section_index:
            section.footer.is_linked_to_previous = False
            section.even_page_footer.is_linked_to_previous = False
        _set_page_number(section.footer, odd=True)
        _set_page_number(section.even_page_footer, odd=False)


def _enable_odd_even_footers(document) -> None:
    settings = document.settings.element
    if settings.find(qn("w:evenAndOddHeaders")) is None:
        settings.append(OxmlElement("w:evenAndOddHeaders"))


def _enable_field_updates(document) -> None:
    settings = document.settings.element
    update_fields = settings.find(qn("w:updateFields"))
    if update_fields is None:
        update_fields = OxmlElement("w:updateFields")
        settings.append(update_fields)
    update_fields.set(qn("w:val"), "true")


def _set_document_grid(section) -> None:
    section_properties = section._sectPr
    for old in section_properties.findall(qn("w:docGrid")):
        section_properties.remove(old)
    grid = OxmlElement("w:docGrid")
    grid.set(qn("w:type"), "linesAndChars")
    grid.set(qn("w:linePitch"), str(EXACT_LINE_SPACING_TWIPS))
    grid.set(qn("w:charSpace"), "0")
    section_properties.append(grid)


def _set_page_number(footer, *, odd: bool) -> None:
    paragraph = _find_page_number_paragraph(footer)
    _clear_paragraph_content(paragraph)
    paragraph.alignment = (
        WD_ALIGN_PARAGRAPH.RIGHT if odd else WD_ALIGN_PARAGRAPH.LEFT
    )
    paragraph_properties = paragraph._p.get_or_add_pPr()
    for old in paragraph_properties.findall(qn("w:ind")):
        paragraph_properties.remove(old)
    side = "right" if odd else "left"
    paragraph_properties.append(
        parse_xml(
            f'<w:ind {nsdecls("w")} w:{side}="{ONE_CHAR_INDENT_TWIPS}" '
            f'w:{side}Chars="{ONE_CHAR_INDENT_CHARS}"/>'
        )
    )

    _append_text_run(paragraph, "－")
    _append_page_field(paragraph)
    _append_text_run(paragraph, "－")


def _find_page_number_paragraph(footer):
    for paragraph in footer.paragraphs:
        if "PAGE" in paragraph._p.xml:
            return paragraph
    for paragraph in footer.paragraphs:
        if not paragraph.text.strip():
            return paragraph
    return footer.add_paragraph()


def _clear_paragraph_content(paragraph) -> None:
    for child in list(paragraph._p):
        if child.tag != qn("w:pPr"):
            paragraph._p.remove(child)


def _append_text_run(paragraph, text: str) -> None:
    run = paragraph.add_run(text)
    _set_page_number_run_font(run)


def _append_page_field(paragraph) -> None:
    begin = paragraph.add_run()
    begin._r.append(_field_char("begin"))
    _set_page_number_run_font(begin)

    instruction = paragraph.add_run()
    instruction_text = OxmlElement("w:instrText")
    instruction_text.set(qn("xml:space"), "preserve")
    instruction_text.text = " PAGE \\* MERGEFORMAT "
    instruction._r.append(instruction_text)
    _set_page_number_run_font(instruction)

    separate = paragraph.add_run()
    separate._r.append(_field_char("separate"))
    _set_page_number_run_font(separate)

    result = paragraph.add_run("1")
    _set_page_number_run_font(result)

    end = paragraph.add_run()
    end._r.append(_field_char("end"))
    _set_page_number_run_font(end)


def _field_char(field_type: str):
    field_char = OxmlElement("w:fldChar")
    field_char.set(qn("w:fldCharType"), field_type)
    return field_char


def _set_page_number_run_font(run) -> None:
    run_properties = run._r.get_or_add_rPr()
    for tag in ("w:rFonts", "w:sz", "w:szCs"):
        for old in run_properties.findall(qn(tag)):
            run_properties.remove(old)
    run_properties.append(
        parse_xml(
            f'<w:rFonts {nsdecls("w")} w:eastAsia="{SONG}" '
            f'w:ascii="{LATIN_FONT}" w:hAnsi="{LATIN_FONT}"/>'
        )
    )
    half_points = PAGE_NUMBER_SIZE_PT * 2
    run_properties.append(
        parse_xml(f'<w:sz {nsdecls("w")} w:val="{half_points}"/>')
    )
    run_properties.append(
        parse_xml(f'<w:szCs {nsdecls("w")} w:val="{half_points}"/>')
    )
