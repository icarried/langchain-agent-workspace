from __future__ import annotations

import hashlib
import os
import zipfile
from pathlib import Path

from docx import Document
from langgraph.graph import END, START, StateGraph

from .compliance import compliance_report, evaluate_compliance
from .formatter import format_docx
from .schemas import (
    DOCX_MIME_TYPE,
    FormattedDocumentResult,
    FormattingState,
)

DEFAULT_MAX_BYTES = 20 * 1024 * 1024


def formatting_max_bytes() -> int:
    value = os.getenv("OFFICIAL_DOCUMENT_FORMATTING_MAX_BYTES", "").strip()
    if not value:
        return DEFAULT_MAX_BYTES
    try:
        parsed = int(value)
    except ValueError as exc:
        raise ValueError("OFFICIAL_DOCUMENT_FORMATTING_MAX_BYTES must be an integer") from exc
    if parsed <= 0:
        raise ValueError("OFFICIAL_DOCUMENT_FORMATTING_MAX_BYTES must be positive")
    return parsed


def _output_filename(original_filename: str) -> str:
    safe_name = Path(original_filename).name
    if Path(safe_name).suffix.lower() not in {".doc", ".docx"}:
        raise ValueError("公文格式化仅支持 DOCX 或 DOC 文件")
    return f"{Path(safe_name).stem}-公文格式化.docx"


def _validate_input(state: FormattingState) -> dict[str, object]:
    source = Path(state["source_path"])
    if source.suffix.lower() != ".docx":
        raise ValueError("公文格式化仅支持 DOCX 文件")
    if not source.is_file():
        raise FileNotFoundError(f"公文文件不存在: {source}")
    size = source.stat().st_size
    if size <= 0:
        raise ValueError("公文 DOCX 为空")
    if size > formatting_max_bytes():
        raise ValueError("公文 DOCX 超过允许大小")
    try:
        Document(source)
    except (OSError, ValueError, zipfile.BadZipFile) as exc:
        raise ValueError("公文文件不是有效 DOCX") from exc
    return {
        "source_size": size,
        "output_filename": _output_filename(state["original_filename"]),
    }


def _format_document(state: FormattingState) -> dict[str, object]:
    if state["dry_run"]:
        return {}
    output = Path(state["output_path"])
    output.parent.mkdir(parents=True, exist_ok=True)
    format_docx(state["source_path"], output)
    return {}


def _build_result(state: FormattingState) -> dict[str, object]:
    if state["dry_run"]:
        content = b""
        digest = ""
        size = 0
        findings = evaluate_compliance(
            Document(state["source_path"]),
            formatted=False,
        )
    else:
        output = Path(state["output_path"])
        try:
            formatted_document = Document(output)
        except (OSError, ValueError, zipfile.BadZipFile) as exc:
            raise ValueError("格式化输出不是有效 DOCX") from exc
        content = output.read_bytes()
        digest = hashlib.sha256(content).hexdigest()
        size = len(content)
        findings = evaluate_compliance(formatted_document, formatted=True)
    report = compliance_report(findings, dry_run=state["dry_run"])
    result = FormattedDocumentResult(
        filename=state["output_filename"],
        mime_type=DOCX_MIME_TYPE,
        content=content,
        sha256=digest,
        size=size,
        dry_run=state["dry_run"],
        report=report,
        findings=findings,
        output_path=state["output_path"] if state.get("persist_output") else "",
    )
    return {"result": result}


def build_graph():
    graph = StateGraph(FormattingState)
    graph.add_node("validate_input", _validate_input)
    graph.add_node("format_document", _format_document)
    graph.add_node("build_result", _build_result)
    graph.add_edge(START, "validate_input")
    graph.add_edge("validate_input", "format_document")
    graph.add_edge("format_document", "build_result")
    graph.add_edge("build_result", END)
    return graph.compile()
