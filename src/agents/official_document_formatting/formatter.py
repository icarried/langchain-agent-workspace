"""Deterministic DOCX formatter for the company's official-document rules.

The formatter changes presentation only. Paragraph text, table cell text, and
their order are treated as immutable input and verified before the result is
published.
"""

from __future__ import annotations

import os
import re
import uuid
from pathlib import Path

from docx import Document
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import parse_xml
from docx.oxml.ns import nsdecls, qn
from docx.shared import Pt

from .attachment_blocks import normalize_attachment_blocks, semantic_paragraph_texts
from .backmatter import apply_backmatter
from .page_layout import apply_page_layout
from .roles import ParagraphRole, classify_paragraphs
from .standards import (
    EXACT_LINE_SPACING_TWIPS,
    FS,
    LATIN_FONT,
    ROLE_STYLES,
    TABLE_SIZE_PT,
)


def set_font(run, name: str, size: int) -> None:
    """Set explicit east-Asian and Latin fonts without requiring server fonts."""
    rpr = run._element.find(qn("w:rPr"))
    if rpr is None:
        rpr = parse_xml(f'<w:rPr {nsdecls("w")}/>')
        run._element.insert(0, rpr)
    for tag in ["w:rFonts", "w:sz", "w:szCs", "w:color", "w:b", "w:bCs"]:
        for old in rpr.findall(qn(tag)):
            rpr.remove(old)
    rpr.append(
        parse_xml(
            f'<w:rFonts {nsdecls("w")} '
            f'w:eastAsia="{name}" w:ascii="{LATIN_FONT}" '
            f'w:hAnsi="{LATIN_FONT}"/>'
        )
    )
    rpr.append(parse_xml(f'<w:sz {nsdecls("w")} w:val="{size * 2}"/>'))
    rpr.append(parse_xml(f'<w:szCs {nsdecls("w")} w:val="{size * 2}"/>'))
    rpr.append(parse_xml(f'<w:color {nsdecls("w")} w:val="000000"/>'))


def document_content_snapshot(document) -> tuple[object, ...]:
    """Return content that formatting is not allowed to change."""
    paragraphs = tuple(paragraph.text for paragraph in document.paragraphs)
    tables = tuple(
        tuple(tuple(cell.text for cell in row.cells) for row in table.rows)
        for table in document.tables
    )
    return paragraphs, tables


def document_semantic_snapshot(document) -> tuple[object, ...]:
    """Return content while ignoring authorized attachment layout whitespace."""
    tables = tuple(
        tuple(tuple(cell.text for cell in row.cells) for row in table.rows)
        for table in document.tables
    )
    return semantic_paragraph_texts(document), tables


def format_docx(src: str | Path, dst: str | Path) -> None:
    """Format one DOCX, allowing only controlled attachment layout normalization."""
    source = Path(src)
    destination = Path(dst)
    document = Document(source)
    source_semantic_snapshot = document_semantic_snapshot(document)

    normalize_attachment_blocks(document)
    normalized_snapshot = document_content_snapshot(document)

    apply_page_layout(document)
    roles = _format_paragraphs(document)
    apply_backmatter(document, roles)
    _format_tables(document)

    destination.parent.mkdir(parents=True, exist_ok=True)
    temporary = destination.with_name(
        f".{destination.name}.{uuid.uuid4().hex}.tmp.docx"
    )
    try:
        document.save(temporary)
        formatted = Document(temporary)
        if document_content_snapshot(formatted) != normalized_snapshot:
            raise ValueError("格式化前后公文正文或表格内容不一致，已拒绝输出")
        if document_semantic_snapshot(formatted) != source_semantic_snapshot:
            raise ValueError("附件换行处理改变了公文实质内容，已拒绝输出")
        os.replace(temporary, destination)
    finally:
        temporary.unlink(missing_ok=True)


def _format_paragraphs(document) -> list[ParagraphRole]:
    roles = classify_paragraphs(
        [paragraph.text for paragraph in document.paragraphs]
    )
    for paragraph, role in zip(document.paragraphs, roles, strict=True):
        if role is ParagraphRole.EMPTY:
            continue
        _reset_paragraph_properties(paragraph)
        spec = ROLE_STYLES[role]
        _apply_paragraph_format(
            paragraph,
            font=spec.font,
            size=spec.size_pt,
            alignment=spec.alignment,
            first_line_indent=spec.first_line_twips,
            first_line_chars=spec.first_line_chars,
        )
    return roles


def _reset_paragraph_properties(paragraph) -> None:
    paragraph.style = paragraph.part.document.styles["Normal"]
    ppr = paragraph._p.get_or_add_pPr()
    for tag in ["w:pBdr", "w:bdr", "w:shd", "w:spacing", "w:ind", "w:jc"]:
        for element in ppr.findall(qn(tag)):
            ppr.remove(element)
    ppr.append(
        parse_xml(
            f'<w:spacing {nsdecls("w")} w:before="0" w:after="0" '
            f'w:line="{EXACT_LINE_SPACING_TWIPS}" w:lineRule="exact"/>'
        )
    )


def _apply_paragraph_format(
    paragraph,
    *,
    font: str,
    size: int,
    alignment,
    first_line_indent: int,
    first_line_chars: int,
) -> None:
    paragraph.alignment = alignment
    ppr = paragraph._p.get_or_add_pPr()
    indentation = (
        f'w:firstLine="{first_line_indent}" w:firstLineChars="{first_line_chars}"'
        if first_line_indent
        else ""
    )
    ppr.append(parse_xml(f'<w:ind {nsdecls("w")} {indentation}/>'))
    for run in paragraph.runs:
        set_font(run, font, size)


def _format_tables(document) -> None:
    for table in document.tables:
        _set_three_line_borders(table)
        _set_table_row_pagination(table)
        alignments = _table_alignments(table)
        for row_index, row in enumerate(table.rows):
            for column_index, cell in enumerate(row.cells):
                cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
                alignment = (
                    WD_ALIGN_PARAGRAPH.CENTER
                    if row_index == 0
                    else alignments[column_index]
                )
                for paragraph in cell.paragraphs:
                    paragraph.alignment = alignment
                    paragraph.paragraph_format.space_before = Pt(2)
                    paragraph.paragraph_format.space_after = Pt(2)
                    for run in paragraph.runs:
                        set_font(run, FS, TABLE_SIZE_PT)
        _set_header_bottom_border(table)


def _set_table_row_pagination(table) -> None:
    if not table.rows:
        return
    for row_index, row in enumerate(table.rows):
        row_properties = row._tr.get_or_add_trPr()
        for tag in ("w:tblHeader", "w:cantSplit"):
            for old in row_properties.findall(qn(tag)):
                row_properties.remove(old)
        if row_index == 0:
            row_properties.append(
                parse_xml(f'<w:tblHeader {nsdecls("w")} w:val="true"/>')
            )
        row_properties.append(
            parse_xml(f'<w:cantSplit {nsdecls("w")} w:val="true"/>')
        )


def _table_alignments(table) -> list[WD_ALIGN_PARAGRAPH]:
    if not table.rows:
        return []
    result = []
    for column_index, header_cell in enumerate(table.rows[0].cells):
        header = header_cell.text.replace("\n", "").strip()
        if any(keyword in header for keyword in ("金额", "价格", "单价", "总价", "预算")):
            result.append(WD_ALIGN_PARAGRAPH.RIGHT)
        elif any(keyword in header for keyword in ("序号", "编号", "单位", "数量", "日期")):
            result.append(WD_ALIGN_PARAGRAPH.CENTER)
        elif any(keyword in header for keyword in ("名称", "规格", "内容", "说明", "备注")):
            result.append(WD_ALIGN_PARAGRAPH.LEFT)
        else:
            values = [row.cells[column_index].text.strip() for row in table.rows[1:]]
            result.append(
                WD_ALIGN_PARAGRAPH.RIGHT
                if values and all(_looks_numeric(value) for value in values if value)
                else WD_ALIGN_PARAGRAPH.CENTER
            )
    return result


def _looks_numeric(value: str) -> bool:
    normalized = value.replace(",", "").replace("，", "").replace("¥", "").strip()
    return bool(normalized) and bool(re.fullmatch(r"[-+]?\d+(?:\.\d+)?%?", normalized))


def _set_three_line_borders(table) -> None:
    tbl_pr = table._tbl.tblPr
    for old_border in tbl_pr.findall(qn("w:tblBorders")):
        tbl_pr.remove(old_border)
    tbl_pr.append(
        parse_xml(
            f'<w:tblBorders {nsdecls("w")}>'
            '  <w:top w:val="single" w:sz="12" w:space="0" w:color="000000"/>'
            '  <w:bottom w:val="single" w:sz="12" w:space="0" w:color="000000"/>'
            '  <w:left w:val="none" w:sz="0" w:space="0" w:color="000000"/>'
            '  <w:right w:val="none" w:sz="0" w:space="0" w:color="000000"/>'
            '  <w:insideH w:val="none" w:sz="0" w:space="0" w:color="000000"/>'
            '  <w:insideV w:val="none" w:sz="0" w:space="0" w:color="000000"/>'
            "</w:tblBorders>"
        )
    )


def _set_header_bottom_border(table) -> None:
    if not table.rows:
        return
    for cell in table.rows[0].cells:
        tc_pr = cell._tc.get_or_add_tcPr()
        for old_border in tc_pr.findall(qn("w:tcBorders")):
            tc_pr.remove(old_border)
        tc_pr.append(
            parse_xml(
                f'<w:tcBorders {nsdecls("w")}>'
                '  <w:bottom w:val="single" w:sz="6" w:space="0" '
                'w:color="000000"/>'
                "</w:tcBorders>"
            )
        )
