"""Attachment, closing, and imprint paragraph-property rules."""

from __future__ import annotations

from docx.oxml import OxmlElement, parse_xml
from docx.oxml.ns import nsdecls, qn

from .roles import ParagraphRole
from .standards import EXACT_LINE_SPACING_TWIPS

_CHAR_TWIPS = 320


def apply_backmatter(document, roles: list[ParagraphRole]) -> None:
    """Apply structural rules that depend on neighboring paragraph roles."""
    imprint_indices = [
        index for index, role in enumerate(roles) if role is ParagraphRole.IMPRINT
    ]
    for index, (paragraph, role) in enumerate(
        zip(document.paragraphs, roles, strict=True)
    ):
        if role is ParagraphRole.ATTACHMENT_NOTE:
            _set_indentation(
                paragraph,
                first_line_twips=2 * _CHAR_TWIPS,
                first_line_chars=200,
            )
            _set_spacing_before(paragraph, _blank_row_before(document, index))
            _set_keep(paragraph, keep_with_next=True, keep_lines=True)
        elif role is ParagraphRole.ATTACHMENT_ITEM:
            _set_indentation(
                paragraph,
                left_twips=2 * _CHAR_TWIPS,
                left_chars=200,
            )
            _set_keep(paragraph, keep_lines=True)
        elif role in {ParagraphRole.SIGNATURE, ParagraphRole.DATE}:
            _set_indentation(paragraph, right_twips=4 * _CHAR_TWIPS, right_chars=400)
            _set_keep(paragraph, keep_lines=True)
        elif role is ParagraphRole.OFFICIAL_ATTACHMENT_MARKER:
            _set_presence_property(paragraph, "w:pageBreakBefore")
        elif role is ParagraphRole.OFFICIAL_ATTACHMENT_TITLE:
            _set_spacing_before(paragraph, _blank_row_before(document, index))
            _set_keep(paragraph, keep_with_next=True, keep_lines=True)
        elif role is ParagraphRole.IMPRINT:
            _set_indentation(
                paragraph,
                left_twips=_CHAR_TWIPS,
                left_chars=100,
                right_twips=_CHAR_TWIPS,
                right_chars=100,
            )
            _set_keep(
                paragraph,
                keep_with_next=index != imprint_indices[-1],
                keep_lines=True,
            )
    if imprint_indices:
        _set_imprint_border(document.paragraphs[imprint_indices[0]], "w:top")
        _set_imprint_border(document.paragraphs[imprint_indices[-1]], "w:bottom")


def _blank_row_before(document, index: int) -> int:
    if index > 0 and not document.paragraphs[index - 1].text.strip():
        return 0
    return EXACT_LINE_SPACING_TWIPS


def _set_spacing_before(paragraph, before_twips: int) -> None:
    paragraph_properties = paragraph._p.get_or_add_pPr()
    spacing = paragraph_properties.find(qn("w:spacing"))
    if spacing is None:
        spacing = OxmlElement("w:spacing")
        paragraph_properties.append(spacing)
    spacing.set(qn("w:before"), str(before_twips))
    spacing.set(qn("w:after"), "0")
    spacing.set(qn("w:line"), str(EXACT_LINE_SPACING_TWIPS))
    spacing.set(qn("w:lineRule"), "exact")


def _set_indentation(
    paragraph,
    *,
    left_twips: int | None = None,
    left_chars: int | None = None,
    right_twips: int | None = None,
    right_chars: int | None = None,
    hanging_twips: int | None = None,
    hanging_chars: int | None = None,
    first_line_twips: int | None = None,
    first_line_chars: int | None = None,
) -> None:
    paragraph_properties = paragraph._p.get_or_add_pPr()
    for old in paragraph_properties.findall(qn("w:ind")):
        paragraph_properties.remove(old)
    values = {
        "left": left_twips,
        "leftChars": left_chars,
        "right": right_twips,
        "rightChars": right_chars,
        "hanging": hanging_twips,
        "hangingChars": hanging_chars,
        "firstLine": first_line_twips,
        "firstLineChars": first_line_chars,
    }
    attributes = " ".join(
        f'w:{name}="{value}"' for name, value in values.items() if value is not None
    )
    paragraph_properties.append(
        parse_xml(f'<w:ind {nsdecls("w")} {attributes}/>')
    )


def _set_keep(
    paragraph,
    *,
    keep_with_next: bool = False,
    keep_lines: bool = False,
) -> None:
    if keep_with_next:
        _set_presence_property(paragraph, "w:keepNext")
    if keep_lines:
        _set_presence_property(paragraph, "w:keepLines")


def _set_presence_property(paragraph, tag: str) -> None:
    paragraph_properties = paragraph._p.get_or_add_pPr()
    if paragraph_properties.find(qn(tag)) is None:
        paragraph_properties.append(OxmlElement(tag))


def _set_imprint_border(paragraph, side: str) -> None:
    paragraph_properties = paragraph._p.get_or_add_pPr()
    borders = paragraph_properties.find(qn("w:pBdr"))
    if borders is None:
        borders = OxmlElement("w:pBdr")
        paragraph_properties.append(borders)
    for old in borders.findall(qn(side)):
        borders.remove(old)
    border = OxmlElement(side)
    border.set(qn("w:val"), "single")
    border.set(qn("w:sz"), "6")
    border.set(qn("w:space"), "1")
    border.set(qn("w:color"), "000000")
    borders.append(border)
