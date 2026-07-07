from __future__ import annotations

from pathlib import Path

from docx import Document
from pypdf import PdfReader

from .schemas import ContractElement


SUPPORTED_EXTENSIONS = {".docx", ".md", ".pdf", ".txt"}


def load_contract_elements(path: str | Path) -> list[ContractElement]:
    contract_path = Path(path)
    if not contract_path.exists():
        raise FileNotFoundError(f"contract file not found: {contract_path}")

    suffix = contract_path.suffix.lower()
    if suffix in {".md", ".txt"}:
        return _load_text(contract_path)
    if suffix == ".docx":
        return _load_docx(contract_path)
    if suffix == ".pdf":
        return _load_pdf(contract_path)

    supported = ", ".join(sorted(SUPPORTED_EXTENSIONS))
    raise ValueError(f"unsupported contract file type '{suffix}', supported: {supported}")


def _load_text(path: Path) -> list[ContractElement]:
    text = path.read_text(encoding="utf-8-sig")
    return _elements_from_lines(text.splitlines(), kind="paragraph", source=path.name)


def _load_docx(path: Path) -> list[ContractElement]:
    document = Document(path)
    elements: list[ContractElement] = []
    for paragraph in document.paragraphs:
        text = paragraph.text.strip()
        if not text:
            continue
        elements.append(
            ContractElement(
                index=len(elements) + 1,
                kind="paragraph",
                text=text,
                source=path.name,
                style=paragraph.style.name if paragraph.style else "",
            )
        )

    for table in document.tables:
        rows = []
        for row in table.rows:
            cells = [cell.text.strip().replace("\n", " ") for cell in row.cells]
            if any(cells):
                rows.append(" | ".join(cells))
        if rows:
            elements.append(
                ContractElement(
                    index=len(elements) + 1,
                    kind="table",
                    text="\n".join(rows),
                    source=path.name,
                )
            )
    return elements


def _load_pdf(path: Path) -> list[ContractElement]:
    reader = PdfReader(str(path))
    elements: list[ContractElement] = []
    for page_number, page in enumerate(reader.pages, start=1):
        page_text = page.extract_text() or ""
        for line in _clean_lines(page_text.splitlines()):
            elements.append(
                ContractElement(
                    index=len(elements) + 1,
                    kind="pdf_line",
                    text=line,
                    source=f"{path.name}:page-{page_number}",
                )
            )
    if not elements:
        raise ValueError("PDF has no extractable text; scanned-image OCR is not supported in v1.")
    return elements


def _elements_from_lines(lines: list[str], *, kind: str, source: str) -> list[ContractElement]:
    return [
        ContractElement(index=index, kind=kind, text=line, source=source)
        for index, line in enumerate(_clean_lines(lines), start=1)
    ]


def _clean_lines(lines: list[str]) -> list[str]:
    return [line.strip() for line in lines if line.strip()]

