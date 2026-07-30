from __future__ import annotations

import mimetypes
import zipfile
from collections.abc import Callable
from contextlib import contextmanager
from contextvars import ContextVar
from pathlib import Path
from typing import Iterator

from src.document_ocr import OCRProvider
from src.knowledge_base.loaders import DocumentRecord

from .settings import DepartmentKnowledgeBaseSettings


TEXT_EXTENSIONS = {".md", ".markdown", ".txt"}
IMAGE_EXTENSIONS = {".bmp", ".jpeg", ".jpg", ".png", ".tif", ".tiff", ".webp"}
SUPPORTED_EXTENSIONS = TEXT_EXTENSIONS | IMAGE_EXTENSIONS | {".docx", ".pdf"}
DocumentProgress = Callable[[str, str], None]
_DOCUMENT_PROGRESS: ContextVar[DocumentProgress | None] = ContextVar(
    "department_kb_document_progress",
    default=None,
)


@contextmanager
def document_progress(callback: DocumentProgress | None) -> Iterator[None]:
    token = _DOCUMENT_PROGRESS.set(callback)
    try:
        yield
    finally:
        _DOCUMENT_PROGRESS.reset(token)


def _emit_progress(stage: str, message: str) -> None:
    callback = _DOCUMENT_PROGRESS.get()
    if callback:
        callback(stage, message)


class AdaptiveDocumentLoader:
    """Prefer local text extraction and call OCR only for image-based content."""

    def __init__(
        self,
        ocr: OCRProvider,
        settings: DepartmentKnowledgeBaseSettings | None = None,
    ) -> None:
        self.ocr = ocr
        self.settings = settings or DepartmentKnowledgeBaseSettings()

    def __call__(self, path: Path, *, source_root: Path) -> DocumentRecord | None:
        suffix = path.suffix.lower()
        if suffix in TEXT_EXTENSIONS:
            text = path.read_text(encoding="utf-8-sig").strip()
        elif suffix == ".pdf":
            text = self._load_pdf(path)
        elif suffix == ".docx":
            text = self._load_docx(path)
        elif suffix in IMAGE_EXTENSIONS:
            _emit_progress("ocr", f"正在 OCR 图片：{path.name}。")
            text = self.ocr.extract_image(
                path.read_bytes(),
                _image_mime_type(path),
                source=path.name,
            ).strip()
            _emit_progress("ocr", f"OCR 已完成：{path.name}。")
        else:
            raise ValueError(f"unsupported knowledge-base document extension: {suffix}")
        if not text:
            return None
        return DocumentRecord(
            text,
            {
                "source": path.relative_to(source_root).as_posix(),
                "file_name": path.name,
                "extension": suffix,
            },
        )

    def _load_pdf(self, path: Path) -> str:
        from pypdf import PdfReader

        reader = PdfReader(str(path))
        if len(reader.pages) > self.settings.ocr_max_pages:
            raise ValueError(
                f"{path.name!r} has {len(reader.pages)} pages; "
                f"limit is {self.settings.ocr_max_pages}"
            )
        rendered = None
        pages: list[str] = []
        for index, page in enumerate(reader.pages):
            local_text = (page.extract_text() or "").strip()
            if len(local_text) >= self.settings.min_local_text_chars:
                text = local_text
            else:
                _emit_progress(
                    "ocr",
                    f"正在 OCR：{path.name}，第 {index + 1}/{len(reader.pages)} 页。",
                )
                if rendered is None:
                    import fitz

                    rendered = fitz.open(path)
                pixmap = rendered[index].get_pixmap(matrix=fitz.Matrix(2, 2), alpha=False)
                text = self.ocr.extract_image(
                    pixmap.tobytes("png"),
                    "image/png",
                    source=f"{path.name}#page-{index + 1}",
                ).strip()
                _emit_progress(
                    "ocr",
                    f"OCR 已完成：{path.name}，第 {index + 1}/{len(reader.pages)} 页。",
                )
            if text:
                pages.append(f"<!-- page:{index + 1} -->\n{text}")
        if rendered is not None:
            rendered.close()
        return "\n\n".join(pages)

    def _load_docx(self, path: Path) -> str:
        from docx import Document

        document = Document(str(path))
        paragraphs = [paragraph.text.strip() for paragraph in document.paragraphs]
        tables = [
            cell.text.strip()
            for table in document.tables
            for row in table.rows
            for cell in row.cells
        ]
        local_text = "\n".join(item for item in paragraphs + tables if item).strip()
        if len(local_text) >= self.settings.min_local_text_chars:
            return local_text

        extracted: list[str] = [local_text] if local_text else []
        with zipfile.ZipFile(path) as archive:
            media = sorted(
                name
                for name in archive.namelist()
                if name.startswith("word/media/")
                and Path(name).suffix.lower() in IMAGE_EXTENSIONS
            )
            if len(media) > self.settings.ocr_max_pages:
                raise ValueError(
                    f"{path.name!r} contains {len(media)} images; "
                    f"limit is {self.settings.ocr_max_pages}"
                )
            for index, name in enumerate(media, start=1):
                image_path = Path(name)
                _emit_progress(
                    "ocr",
                    f"正在 OCR：{path.name}，第 {index}/{len(media)} 个内嵌图片。",
                )
                extracted.append(
                    self.ocr.extract_image(
                        archive.read(name),
                        _image_mime_type(image_path),
                        source=f"{path.name}#image-{index}",
                    ).strip()
                )
                _emit_progress(
                    "ocr",
                    f"OCR 已完成：{path.name}，第 {index}/{len(media)} 个内嵌图片。",
                )
        return "\n\n".join(item for item in extracted if item)


def _image_mime_type(path: Path) -> str:
    return mimetypes.guess_type(path.name)[0] or "image/png"
