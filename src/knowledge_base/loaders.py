from __future__ import annotations

from dataclasses import dataclass, field
from pathlib import Path


SUPPORTED_EXTENSIONS = {".pdf", ".docx", ".md", ".markdown", ".txt"}


@dataclass(slots=True)
class DocumentRecord:
    page_content: str
    metadata: dict[str, object] = field(default_factory=dict)


def iter_supported_paths(
    root: Path,
    *,
    supported_extensions: set[str] | None = None,
) -> list[Path]:
    if not root.exists():
        return []
    extensions = supported_extensions or SUPPORTED_EXTENSIONS
    return sorted(
        path
        for path in root.rglob("*")
        if path.is_file() and path.suffix.lower() in extensions
    )


def load_document(path: Path, *, source_root: Path) -> DocumentRecord | None:
    suffix = path.suffix.lower()
    if suffix == ".pdf":
        from pypdf import PdfReader

        reader = PdfReader(str(path))
        text = "\n\n".join((page.extract_text() or "").strip() for page in reader.pages)
    elif suffix == ".docx":
        from docx import Document

        document = Document(str(path))
        paragraphs = [paragraph.text.strip() for paragraph in document.paragraphs]
        tables = [cell.text.strip() for table in document.tables for row in table.rows for cell in row.cells]
        text = "\n".join(item for item in paragraphs + tables if item)
    else:
        text = path.read_text(encoding="utf-8")
    text = text.strip()
    if not text:
        return None
    source = path.relative_to(source_root).as_posix()
    return DocumentRecord(text, {"source": source, "file_name": path.name, "extension": suffix})


def chunk_documents(
    documents: list[DocumentRecord], *, chunk_size: int = 800, chunk_overlap: int = 100
) -> list[DocumentRecord]:
    if chunk_size <= 0 or chunk_overlap < 0 or chunk_overlap >= chunk_size:
        raise ValueError("invalid chunk size or overlap")
    result: list[DocumentRecord] = []
    step = chunk_size - chunk_overlap
    for document in documents:
        for index, start in enumerate(range(0, len(document.page_content), step)):
            text = document.page_content[start : start + chunk_size].strip()
            if text:
                metadata = dict(document.metadata)
                metadata["chunk_index"] = index
                result.append(DocumentRecord(text, metadata))
            if start + chunk_size >= len(document.page_content):
                break
    return result
