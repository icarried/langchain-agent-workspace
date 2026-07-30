from __future__ import annotations

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn

from src.agents.official_document_formatting.page_layout import apply_page_layout


def test_applies_page_grid_and_odd_even_page_numbers() -> None:
    document = Document()
    document.add_paragraph("测试正文")

    apply_page_layout(document)

    section = document.sections[0]
    assert round(section.page_width.cm, 1) == 21.0
    assert round(section.page_height.cm, 1) == 29.7
    assert round(section.top_margin.cm, 1) == 3.7
    assert round(section.bottom_margin.cm, 1) == 3.5
    assert round(section.left_margin.cm, 1) == 2.8
    assert round(section.right_margin.cm, 1) == 2.6
    assert round(section.footer_distance.cm, 1) == 2.5

    settings = document.settings.element
    assert settings.find(qn("w:evenAndOddHeaders")) is not None
    update_fields = settings.find(qn("w:updateFields"))
    assert update_fields is not None
    assert update_fields.get(qn("w:val")) == "true"

    grid = section._sectPr.find(qn("w:docGrid"))
    assert grid is not None
    assert grid.get(qn("w:type")) == "linesAndChars"
    assert grid.get(qn("w:linePitch")) == "560"
    assert grid.get(qn("w:charSpace")) == "0"

    odd = section.footer.paragraphs[0]
    even = section.even_page_footer.paragraphs[0]
    assert odd.alignment == WD_ALIGN_PARAGRAPH.RIGHT
    assert even.alignment == WD_ALIGN_PARAGRAPH.LEFT
    assert "PAGE" in odd._p.xml
    assert "PAGE" in even._p.xml
    assert odd.text == "－1－"
    assert even.text == "－1－"
    assert odd._p.pPr.ind.get(qn("w:rightChars")) == "100"
    assert even._p.pPr.ind.get(qn("w:leftChars")) == "100"

    for paragraph in (odd, even):
        for run in paragraph.runs:
            run_properties = run._r.get_or_add_rPr()
            fonts = run_properties.rFonts
            assert fonts.get(qn("w:eastAsia")) == "宋体"
            assert run_properties.sz.get(qn("w:val")) == "28"


def test_page_number_generation_is_idempotent() -> None:
    document = Document()

    apply_page_layout(document)
    apply_page_layout(document)

    section = document.sections[0]
    assert section.footer._element.xml.count("PAGE") == 1
    assert section.even_page_footer._element.xml.count("PAGE") == 1
