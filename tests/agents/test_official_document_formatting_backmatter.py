from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn

from src.agents.official_document_formatting.formatter import format_docx


def _paragraph_by_text(document, text: str):
    return next(paragraph for paragraph in document.paragraphs if paragraph.text == text)


def test_formats_attachment_closing_formal_attachment_and_imprint(
    tmp_path: Path,
) -> None:
    source = tmp_path / "input.docx"
    output = tmp_path / "output.docx"
    document = Document()
    document.add_paragraph("关于开展测试工作的通知")
    document.add_paragraph("公司各部门：")
    document.add_paragraph("请按要求执行。")
    document.add_paragraph("附件：1. 测试材料清单")
    document.add_paragraph("2. 测试报价单")
    document.add_paragraph("技术支撑部")
    document.add_paragraph("2026年7月29日")
    document.add_paragraph("（联系人：张三）")
    document.add_paragraph("附件 1")
    document.add_paragraph("测试材料清单")
    document.add_paragraph("一、材料内容")
    document.add_paragraph("抄送：公司领导。")
    document.add_paragraph("综合管理部 2026年7月29日印发")
    document.save(source)

    format_docx(source, output)

    formatted = Document(output)
    note = _paragraph_by_text(formatted, "附件：")
    note_ind = note._p.pPr.ind
    assert note_ind.get(qn("w:firstLineChars")) == "200"
    assert note._p.pPr.spacing.get(qn("w:before")) == "560"
    assert note._p.pPr.keepNext is not None

    first_item = _paragraph_by_text(formatted, "1. 测试材料清单")
    second_item = _paragraph_by_text(formatted, "2. 测试报价单")
    for item in (first_item, second_item):
        assert item._p.pPr.ind.get(qn("w:leftChars")) == "200"
        assert item.text == item.text.lstrip()

    signature = _paragraph_by_text(formatted, "技术支撑部")
    date = _paragraph_by_text(formatted, "2026年7月29日")
    assert signature.alignment == WD_ALIGN_PARAGRAPH.RIGHT
    assert date.alignment == WD_ALIGN_PARAGRAPH.RIGHT
    assert signature._p.pPr.ind.get(qn("w:rightChars")) == "400"
    assert date._p.pPr.ind.get(qn("w:rightChars")) == "400"

    marker = _paragraph_by_text(formatted, "附件 1")
    attachment_title = _paragraph_by_text(formatted, "测试材料清单")
    assert marker._p.pPr.pageBreakBefore is not None
    assert attachment_title.alignment == WD_ALIGN_PARAGRAPH.CENTER
    assert attachment_title._p.pPr.spacing.get(qn("w:before")) == "560"

    copy_line = _paragraph_by_text(formatted, "抄送：公司领导。")
    issue_line = _paragraph_by_text(formatted, "综合管理部 2026年7月29日印发")
    for paragraph in (copy_line, issue_line):
        indentation = paragraph._p.pPr.ind
        assert indentation.get(qn("w:leftChars")) == "100"
        assert indentation.get(qn("w:rightChars")) == "100"
        run_properties = paragraph.runs[0]._r.get_or_add_rPr()
        assert run_properties.sz.get(qn("w:val")) == "28"
        assert paragraph._p.pPr.keepLines is not None
    assert copy_line._p.pPr.keepNext is not None
    copy_borders = copy_line._p.pPr.find(qn("w:pBdr"))
    issue_borders = issue_line._p.pPr.find(qn("w:pBdr"))
    assert copy_borders.find(qn("w:top")) is not None
    assert issue_borders.find(qn("w:bottom")) is not None


def test_splits_ascii_attachment_label_and_removes_item_layout_spaces(
    tmp_path: Path,
) -> None:
    source = tmp_path / "input.docx"
    output = tmp_path / "output.docx"
    document = Document()
    document.add_paragraph("关于开展测试工作的通知")
    document.add_paragraph("请按要求执行。")
    document.add_paragraph("附件: 1. 测试材料清单")
    document.add_paragraph("      2. 测试报价单")
    document.add_paragraph("技术支撑部")
    document.add_paragraph("2026年7月29日")
    document.save(source)

    format_docx(source, output)

    formatted = Document(output)
    attachment_texts = [
        paragraph.text
        for paragraph in formatted.paragraphs
        if paragraph.text.startswith(("附件", "1.", "2."))
    ]
    assert attachment_texts == ["附件:", "1. 测试材料清单", "2. 测试报价单"]
