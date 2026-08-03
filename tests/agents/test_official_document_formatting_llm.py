from __future__ import annotations

import base64
import hashlib
import json
from pathlib import Path

from docx import Document
from fastapi.testclient import TestClient
from fastmcp import Client
import pytest

from src.agents.official_document_formatting.openai_compatible_api import (
    MODEL_ID,
    ChatCompletionRequest,
    app,
    parse_document_request,
)
from src.agents.official_document_formatting.mcp_server import mcp
import src.agents.official_document_formatting.service as formatting_service
import src.agents.official_document_formatting.mcp_server as formatting_mcp_server


def _make_docx(path: Path) -> Path:
    document = Document()
    document.add_paragraph("关于开展测试工作的请示")
    document.add_paragraph("请按要求办理。")
    document.save(path)
    return path


def _payload(path: Path, *, stream: bool, dry_run: bool = False) -> dict[str, object]:
    return {
        "model": MODEL_ID,
        "stream": stream,
        "dry_run": dry_run,
        "messages": [
            {
                "role": "user",
                "content": f"公文文件：\n{path}\n\n输出要求：请格式化公文。",
            }
        ],
    }


def test_models_and_readiness_probe() -> None:
    client = TestClient(app)
    health = client.get("/health")
    assert health.status_code == 200
    assert health.json() == {
        "status": "ok",
        "agent": "official-document-formatting",
        "model": MODEL_ID,
    }

    models = client.get("/v1/models")
    assert models.status_code == 200
    assert models.json()["data"][0]["id"] == MODEL_ID

    response = client.post(
        "/v1/chat/completions",
        json={"model": MODEL_ID, "messages": [{"role": "user", "content": "hello"}]},
    )
    assert response.status_code == 200
    assert "公文文件" in response.json()["choices"][0]["message"]["content"]


def test_parse_request_accepts_platform_attachment_and_file_url() -> None:
    attachment_request = ChatCompletionRequest(
        model=MODEL_ID,
        messages=[
            {
                "role": "user",
                "content": "附件：\n- 待格式化公文.docx: http://minio.example/input.docx?signature=abc",
            }
        ],
    )
    assert parse_document_request(attachment_request).document_path.startswith(
        "http://minio.example/input.docx"
    )

    part_request = ChatCompletionRequest(
        model=MODEL_ID,
        messages=[
            {
                "role": "user",
                "content": [
                    {"type": "text", "text": "请格式化"},
                    {
                        "type": "file_url",
                        "file_url": {"url": "http://minio.example/input.docx"},
                    },
                ],
            }
        ],
    )
    assert parse_document_request(part_request).document_path.endswith("input.docx")


def test_non_stream_completion_returns_file_payload(tmp_path: Path) -> None:
    source = _make_docx(tmp_path / "input.docx")

    response = TestClient(app).post(
        "/v1/chat/completions", json=_payload(source, stream=False)
    )

    assert response.status_code == 200
    message = response.json()["choices"][0]["message"]
    file_payload = message["file"]
    content = base64.b64decode(file_payload["content_base64"], validate=True)
    assert file_payload["status"] == "completed"
    assert file_payload["file_type"] == "docx"
    assert file_payload["filename"] == "input-公文格式化.docx"
    assert hashlib.sha256(content).hexdigest() == file_payload["sha256"]
    assert content.startswith(b"PK")
    assert "未验证" in message["content"]


def test_stream_completion_returns_delta_file_and_done(tmp_path: Path) -> None:
    source = _make_docx(tmp_path / "input.docx")

    with TestClient(app).stream(
        "POST", "/v1/chat/completions", json=_payload(source, stream=True)
    ) as response:
        text = "".join(response.iter_text())

    assert response.status_code == 200
    assert "reasoning_content" in text
    assert "data: [DONE]" in text
    chunks = [
        json.loads(line.removeprefix("data: "))
        for line in text.splitlines()
        if line.startswith("data: {")
    ]
    file_chunks = [
        chunk["choices"][0]["delta"]["file"]
        for chunk in chunks
        if "file" in chunk["choices"][0]["delta"]
    ]
    assert len(file_chunks) == 1
    assert file_chunks[0]["file_type"] == "docx"


def test_dry_run_returns_report_without_file(tmp_path: Path) -> None:
    source = _make_docx(tmp_path / "input.docx")

    response = TestClient(app).post(
        "/v1/chat/completions", json=_payload(source, stream=False, dry_run=True)
    )

    assert response.status_code == 200
    message = response.json()["choices"][0]["message"]
    assert "dry-run" in message["content"]
    assert "未验证" in message["content"]
    assert "file" not in message


def test_non_stream_completion_converts_legacy_doc(tmp_path: Path, monkeypatch) -> None:
    converted = _make_docx(tmp_path / "converted.docx").read_bytes()
    source = tmp_path / "input.doc"
    source.write_bytes(b"legacy-doc-placeholder")
    monkeypatch.setattr(
        formatting_service,
        "convert_doc_to_docx",
        lambda data, *, source: converted,
    )

    response = TestClient(app).post(
        "/v1/chat/completions", json=_payload(source, stream=False)
    )

    assert response.status_code == 200
    file_payload = response.json()["choices"][0]["message"]["file"]
    assert file_payload["filename"] == "input-公文格式化.docx"
    assert base64.b64decode(file_payload["content_base64"], validate=True).startswith(b"PK")


@pytest.mark.asyncio
async def test_mcp_formats_docx_and_returns_file(tmp_path: Path) -> None:
    source = _make_docx(tmp_path / "input.docx")
    async with Client(mcp) as client:
        result = await client.call_tool(
            "format_document",
            {
                "document": {
                    "filename": source.name,
                    "content_base64": base64.b64encode(source.read_bytes()).decode(),
                },
                "dry_run": False,
            },
        )

    output = base64.b64decode(result.data["content_base64"], validate=True)
    assert result.data["filename"] == "input-公文格式化.docx"
    assert hashlib.sha256(output).hexdigest() == result.data["sha256"]
    assert output.startswith(b"PK")


@pytest.mark.asyncio
async def test_mcp_formats_legacy_doc_and_returns_docx(
    tmp_path: Path,
    monkeypatch,
) -> None:
    converted = _make_docx(tmp_path / "converted.docx").read_bytes()
    source = tmp_path / "input.doc"
    source.write_bytes(b"legacy-doc-placeholder")
    monkeypatch.setattr(
        formatting_service,
        "convert_doc_to_docx",
        lambda data, *, source: converted,
    )

    async with Client(mcp) as client:
        result = await client.call_tool(
            "format_document",
            {
                "document": {
                    "filename": source.name,
                    "content_base64": base64.b64encode(source.read_bytes()).decode(),
                }
            },
        )

    output = base64.b64decode(result.data["content_base64"], validate=True)
    assert result.data["filename"] == "input-公文格式化.docx"
    assert result.data["mime_type"].endswith("wordprocessingml.document")
    assert output.startswith(b"PK")


@pytest.mark.asyncio
async def test_mcp_formats_minio_url_and_returns_docx(
    tmp_path: Path,
    monkeypatch,
) -> None:
    source = _make_docx(tmp_path / "input.docx")
    minio_url = "http://minio.example:9000/private/input.docx?X-Amz-Signature=abc"
    monkeypatch.setattr(
        formatting_mcp_server,
        "read_remote_file",
        lambda url, *, max_bytes: (source.read_bytes(), "application/vnd.openxmlformats-officedocument.wordprocessingml.document"),
    )

    async with Client(mcp) as client:
        result = await client.call_tool(
            "format_document",
            {"document": {"url": minio_url}},
        )

    output = base64.b64decode(result.data["content_base64"], validate=True)
    assert result.data["filename"] == "input-公文格式化.docx"
    assert output.startswith(b"PK")
