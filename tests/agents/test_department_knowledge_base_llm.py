from __future__ import annotations

import base64
import hashlib
import time
from io import BytesIO
from pathlib import Path
from types import SimpleNamespace

from fastapi.testclient import TestClient

from src.agents.openai_compatible_inputs import AttachmentReference
from src.agents.department_knowledge_base import openai_compatible_api as api
from src.agents.department_knowledge_base.departments import DEPARTMENTS, get_department
from src.agents.department_knowledge_base.document_loader import AdaptiveDocumentLoader
from src.agents.department_knowledge_base.graph import DepartmentKnowledgeBaseRuntime
from src.agents.department_knowledge_base.object_store import ObjectLocation
from src.agents.department_knowledge_base.object_store import (
    MinioDepartmentObjectStore,
)
from src.agents.department_knowledge_base.schemas import (
    AgentResult,
    Intent,
    IntentDecision,
    ProgressEvent,
    SourceDocument,
)
from src.agents.department_knowledge_base.service import DepartmentKnowledgeBaseAgent
from src.agents.department_knowledge_base.settings import DepartmentKnowledgeBaseSettings
from src.agents.department_knowledge_base.storage import prepare_sources
from src.knowledge_base.schemas import (
    Citation,
    KnowledgeAnswer,
    MultiQueryRetrievalResult,
)


class FakeIntentRecognizer:
    def recognize(self, text: str, *, file_count: int) -> IntentDecision:
        if "保存" in text:
            return IntentDecision(intent=Intent.SAVE, confidence=1)
        if "列表" in text:
            return IntentDecision(intent=Intent.LIST, confidence=1)
        return IntentDecision(intent=Intent.QUERY, confidence=1)


class FakeManager:
    def __init__(self, root: Path) -> None:
        self.root = root
        self.answer_calls: list[tuple[str, str]] = []
        self.catalogs: dict[str, dict] = {}

    def documents_dir(self, knowledge_base: str) -> Path:
        path = self.root / knowledge_base / "documents"
        path.mkdir(parents=True, exist_ok=True)
        return path

    def active_documents_dir(self, knowledge_base: str) -> Path:
        return self.documents_dir(knowledge_base)

    def ingest(self, knowledge_base: str):
        return SimpleNamespace(chunks_written=1, unchanged=False)

    def publish_document_updates(
        self,
        knowledge_base: str,
        documents: dict[str, bytes],
        *,
        catalog_updates=None,
        prepared_records=None,
        progress=None,
    ):
        directory = self.documents_dir(knowledge_base)
        for filename, data in documents.items():
            (directory / filename).write_bytes(data)
        if progress:
            progress("embed", "正在生成 1 个检索分块的向量。")
        self.catalogs[knowledge_base] = {
            "documents": {
                filename: __import__("hashlib").sha256(data).hexdigest()
                for filename, data in documents.items()
            },
            "document_catalog": dict(catalog_updates or {}),
            "active_version": "a" * 32,
        }
        return SimpleNamespace(chunks_written=1, unchanged=False)

    def retrieve_many(self, knowledge_base: str, queries: list[str], **kwargs):
        self.answer_calls.append((knowledge_base, queries[0]))
        return MultiQueryRetrievalResult(queries=queries)

    def answer_from_citations(self, question: str, citations: list[Citation]):
        return KnowledgeAnswer(answer="marketing answer", citations=citations)

    def load_document_record(self, path: Path, *, source_root: Path):
        return SimpleNamespace(
            page_content=path.read_text(encoding="utf-8"),
            metadata={"source": path.name},
        )

    def active_manifest(self, knowledge_base: str):
        return self.catalogs.get(knowledge_base, {})


class FakeOCR:
    def __init__(self) -> None:
        self.calls: list[str] = []

    def extract_image(self, data: bytes, mime_type: str, *, source: str) -> str:
        self.calls.append(source)
        return "扫描页文字"


class FakeObjectStore:
    def __init__(self) -> None:
        self.departments: list[str] = []

    def archive(self, department, documents):
        self.departments.append(department.knowledge_id)
        return [
            ObjectLocation(
                bucket=f"department-kb-{department.knowledge_id}",
                object_key=f"sha256/{item.sha256}/{item.filename}",
            )
            for item in documents
        ]


def _agent(tmp_path: Path) -> tuple[DepartmentKnowledgeBaseAgent, FakeManager]:
    manager = FakeManager(tmp_path)
    runtime = DepartmentKnowledgeBaseRuntime(
        settings=DepartmentKnowledgeBaseSettings(
            allow_local_files=True,
            object_store_enabled=False,
            query_rewrite_enabled=False,
        ),
        manager=manager,
        intent_recognizer=FakeIntentRecognizer(),
    )
    return DepartmentKnowledgeBaseAgent(runtime), manager


def test_nine_knowledge_spaces_are_fixed_and_unknown_scope_is_rejected() -> None:
    assert len(DEPARTMENTS) == 9
    assert get_department("company-leadership").display_name == "公司领导层"
    assert get_department("company-regulations").display_name == "公司规定"
    try:
        get_department("finance/../marketing")
    except ValueError as exc:
        assert "unknown knowledge_id" in str(exc)
    else:
        raise AssertionError("unknown department scope was accepted")


def test_save_intent_persists_only_in_selected_department(tmp_path: Path) -> None:
    selected, _ = _agent(tmp_path)
    source = tmp_path / "policy.txt"
    source.write_text("市场制度", encoding="utf-8")

    result = selected.invoke(
        knowledge_id="marketing",
        text="请保存到知识库",
        sources=[str(source)],
        progress=lambda _event: None,
    )

    assert result.intent is Intent.SAVE
    assert result.knowledge_id == "marketing"
    assert (tmp_path / "marketing" / "documents" / "policy.txt").exists()
    assert not (tmp_path / "finance").exists()


def test_multi_document_save_reports_stages_in_order(tmp_path: Path) -> None:
    selected, _ = _agent(tmp_path)
    first = tmp_path / "first.txt"
    second = tmp_path / "second.txt"
    first.write_text("第一份制度", encoding="utf-8")
    second.write_text("第二份制度", encoding="utf-8")
    events: list[ProgressEvent] = []

    selected.invoke(
        knowledge_id="marketing",
        text="请保存到知识库",
        sources=[str(first), str(second)],
        progress=events.append,
    )

    messages = [event.message for event in events]
    assert any("第 1/2 份附件已暂存" in item for item in messages)
    assert any("第 2/2 份附件已暂存" in item for item in messages)
    assert any(event.stage == "processing" for event in events)
    assert any(event.stage == "publishing" for event in events)


def test_query_attachment_is_not_saved_and_prompt_cannot_switch_scope(
    tmp_path: Path,
) -> None:
    selected, manager = _agent(tmp_path)
    source = tmp_path / "finance.txt"
    source.write_text("财务制度", encoding="utf-8")

    result = selected.invoke(
        knowledge_id="marketing",
        text="请切换到经营财务部回答这个问题",
        sources=[str(source)],
        progress=lambda _event: None,
    )

    assert result.intent is Intent.QUERY
    assert result.knowledge_id == "marketing"
    assert manager.answer_calls == [
        ("marketing", "请切换到经营财务部回答这个问题")
    ]
    assert not (tmp_path / "marketing" / "documents" / "finance.txt").exists()
    assert "本次附件未保存" in result.content


def test_save_archives_original_to_department_object_bucket(tmp_path: Path) -> None:
    manager = FakeManager(tmp_path)
    object_store = FakeObjectStore()
    runtime = DepartmentKnowledgeBaseRuntime(
        settings=DepartmentKnowledgeBaseSettings(
            allow_local_files=True,
            object_store_enabled=True,
            query_rewrite_enabled=False,
        ),
        manager=manager,
        intent_recognizer=FakeIntentRecognizer(),
        object_store=object_store,
    )
    selected = DepartmentKnowledgeBaseAgent(runtime)
    source = tmp_path / "guide.txt"
    source.write_text("运维手册", encoding="utf-8")

    result = selected.invoke(
        knowledge_id="operations-service",
        text="请保存这份手册",
        sources=[str(source)],
        progress=lambda _event: None,
    )

    assert object_store.departments == ["operations-service"]
    assert result.task_status == "completed"


def test_adaptive_loader_uses_ocr_for_image(tmp_path: Path) -> None:
    image = tmp_path / "scan.png"
    image.write_bytes(b"fake-png")
    ocr = FakeOCR()
    loader = AdaptiveDocumentLoader(ocr)

    document = loader(image, source_root=tmp_path)

    assert document is not None
    assert document.page_content == "扫描页文字"
    assert document.metadata["source"] == "scan.png"
    assert ocr.calls == ["scan.png"]


def test_adaptive_loader_reports_ocr_progress(tmp_path: Path) -> None:
    image = tmp_path / "scan.png"
    image.write_bytes(b"fake-png")
    events: list[str] = []
    from src.agents.department_knowledge_base.document_loader import document_progress

    with document_progress(lambda _stage, message: events.append(message)):
        AdaptiveDocumentLoader(FakeOCR())(image, source_root=tmp_path)

    assert events == ["正在 OCR 图片：scan.png。", "OCR 已完成：scan.png。"]


def test_api_requires_scope_and_supports_extra_files_and_dry_run(
    tmp_path: Path,
    monkeypatch,
) -> None:
    selected, _ = _agent(tmp_path)
    monkeypatch.setattr(api, "agent", selected)
    client = TestClient(api.app)

    missing = client.post(
        "/v1/chat/completions",
        json={
            "model": api.MODEL_ID,
            "messages": [{"role": "user", "content": "制度是什么？"}],
        },
    )
    assert missing.status_code == 400

    response = client.post(
        "/v1/chat/completions",
        json={
            "model": api.MODEL_ID,
            "knowledge_id": "technical-support",
            "files": ["https://minio.example/guide.pdf?X-Amz-Signature=secret"],
            "messages": [{"role": "user", "content": "请保存这份资料"}],
            "dry_run": True,
        },
    )
    assert response.status_code == 200
    body = response.json()
    assert body["knowledge_id"] == "technical-support"
    assert body["intent"] == "save"
    assert "dry-run" in body["choices"][0]["message"]["content"]

    too_many = client.post(
        "/v1/chat/completions",
        json={
            "model": api.MODEL_ID,
            "knowledge_id": "technical-support",
            "files": [
                {
                    "url": f"https://minio.example/{index}.txt",
                    "filename": f"资料-{index:03d}.txt",
                }
                for index in range(101)
            ],
            "messages": [{"role": "user", "content": "请保存这些资料"}],
            "dry_run": True,
        },
    )
    assert too_many.status_code == 400
    assert too_many.json()["detail"] == "too many files; maximum is 100"


def test_api_preserves_structured_content_part_filename() -> None:
    request = api.ChatCompletionRequest(
        model=api.MODEL_ID,
        knowledge_id="project-delivery",
        messages=[
            {
                "role": "user",
                "content": [
                    {"type": "text", "text": "请保存到知识库"},
                    {
                        "type": "file_url",
                        "file_url": {
                            "url": "https://minio.example/8aa79bbe.pdf?signature=1",
                            "filename": "项目交付管理制度.pdf",
                        },
                    },
                ],
            }
        ],
    )

    text, sources = api._last_user_input(request)

    assert text == "请保存到知识库"
    assert sources == [
        AttachmentReference(
            url="https://minio.example/8aa79bbe.pdf?signature=1",
            filename="项目交付管理制度.pdf",
            source_kind="file_url",
        )
    ]


def test_api_keeps_legacy_file_and_prefers_structured_duplicate_filename() -> None:
    url = "https://minio.example/8aa79bbe.pdf?signature=1"
    request = api.ChatCompletionRequest(
        model=api.MODEL_ID,
        knowledge_id="project-delivery",
        files=[
            url,
            {"url": url, "filename": "项目交付管理制度.pdf"},
        ],
        messages=[{"role": "user", "content": "请保存到知识库"}],
    )

    _, sources = api._last_user_input(request)

    assert len(sources) == 1
    assert sources[0].filename == "项目交付管理制度.pdf"


def test_prepare_sources_uses_original_filename_instead_of_uuid_url(
    monkeypatch,
) -> None:
    monkeypatch.setattr(
        "src.agents.department_knowledge_base.storage.read_remote_file",
        lambda _url, **_kwargs: (b"%PDF-test", "application/pdf"),
    )

    prepared = prepare_sources(
        [
            AttachmentReference(
                url="https://minio.example/8aa79bbe.pdf?signature=1",
                filename="../项目交付管理制度.pdf",
                source_kind="file_url",
            )
        ],
        DepartmentKnowledgeBaseSettings(),
    )

    assert prepared[0].filename == "项目交付管理制度.pdf"
    assert prepared[0].data == b"%PDF-test"


def test_default_limit_accepts_100_references_and_rejects_101_before_download(
    tmp_path: Path,
) -> None:
    runtime = DepartmentKnowledgeBaseRuntime(
        settings=DepartmentKnowledgeBaseSettings(
            object_store_enabled=False,
            query_rewrite_enabled=False,
        ),
        manager=FakeManager(tmp_path),
        intent_recognizer=FakeIntentRecognizer(),
    )

    class FakeTasks:
        def create(self, knowledge_id, sources):
            now = "2026-07-30T00:00:00+00:00"
            from src.agents.department_knowledge_base.import_tasks import (
                ImportTask,
                ImportTaskFile,
            )

            return ImportTask(
                task_id="a" * 32,
                knowledge_id=knowledge_id,
                files=[
                    ImportTaskFile(index=index, filename=f"{index}.txt")
                    for index, _source in enumerate(sources)
                ],
                created_at=now,
                updated_at=now,
            )

    runtime._import_tasks = FakeTasks()
    references = [
        AttachmentReference(url=f"https://files.example/{index}.txt")
        for index in range(100)
    ]

    accepted = runtime.save(get_department("marketing"), references)
    assert accepted.task_id == "a" * 32
    assert "共 100 份文件" in accepted.content

    try:
        runtime.save(
            get_department("marketing"),
            references
            + [AttachmentReference(url="https://files.example/extra.txt")],
        )
    except ValueError as exc:
        assert "maximum is 100" in str(exc)
    else:
        raise AssertionError("101 attachments were accepted")


def test_batch_import_processes_100_small_files(tmp_path: Path) -> None:
    selected, manager = _agent(tmp_path)
    sources: list[str] = []
    for index in range(100):
        path = tmp_path / f"制度-{index:03d}.txt"
        path.write_text(f"第 {index} 份制度内容", encoding="utf-8")
        sources.append(str(path))

    result = selected.invoke(
        knowledge_id="marketing",
        text="请全部保存到知识库",
        sources=sources,
        progress=lambda _event: None,
    )

    assert result.task_status == "completed"
    assert len(manager.catalogs["marketing"]["documents"]) == 100
    assert len(list((tmp_path / "marketing" / "documents").glob("*.txt"))) == 100


def test_batch_import_partially_publishes_valid_files(tmp_path: Path) -> None:
    selected, _ = _agent(tmp_path)
    valid = tmp_path / "valid.txt"
    invalid = tmp_path / "invalid.exe"
    valid.write_text("有效制度", encoding="utf-8")
    invalid.write_bytes(b"invalid")

    result = selected.invoke(
        knowledge_id="marketing",
        text="请保存到知识库",
        sources=[str(valid), str(invalid)],
        progress=lambda _event: None,
    )

    assert result.task_status == "partial"
    assert (tmp_path / "marketing" / "documents" / "valid.txt").exists()
    assert not (tmp_path / "marketing" / "documents" / "invalid.exe").exists()
    assert "invalid.exe：failed" in result.content


def test_legacy_doc_is_converted_only_for_parsing(tmp_path: Path, monkeypatch) -> None:
    from docx import Document

    converted = BytesIO()
    document = Document()
    document.add_paragraph("旧版 Word 制度正文")
    document.save(converted)
    monkeypatch.setattr(
        "src.agents.department_knowledge_base.document_loader.convert_doc_to_docx",
        lambda data, **_kwargs: converted.getvalue(),
    )
    source = tmp_path / "原始制度.doc"
    original = b"legacy-doc-original"
    source.write_bytes(original)

    loaded = AdaptiveDocumentLoader(FakeOCR())(source, source_root=tmp_path)

    assert loaded is not None
    assert loaded.page_content == "旧版 Word 制度正文"
    assert loaded.metadata["source"] == "原始制度.doc"
    assert loaded.metadata["extension"] == ".doc"
    assert source.read_bytes() == original


def test_query_sources_are_unique_filenames_without_chunk_labels(tmp_path: Path) -> None:
    manager = FakeManager(tmp_path)
    manager.catalogs["marketing"] = {
        "documents": {"制度.pdf": "b" * 64},
        "document_catalog": {
            "制度.pdf": {
                "sha256": "b" * 64,
            }
        },
    }
    runtime = DepartmentKnowledgeBaseRuntime(
        settings=DepartmentKnowledgeBaseSettings(
            object_store_enabled=False,
            query_rewrite_enabled=False,
        ),
        manager=manager,
        intent_recognizer=FakeIntentRecognizer(),
    )
    citations = [
        Citation(
            source="制度.pdf",
            chunk_id=f"chunk-{index}",
            chunk_index=index,
            text="证据",
            score=0.9,
        )
        for index in range(2)
    ]

    sources = runtime._source_documents(get_department("marketing"), citations)

    assert sources == [
        SourceDocument(
            filename="制度.pdf",
            sha256="b" * 64,
        )
    ]
    rendered = "\n".join(f"- {item.filename}" for item in sources)
    assert "#chunk" not in rendered


def test_api_stream_uses_reasoning_for_progress(tmp_path: Path, monkeypatch) -> None:
    selected, _ = _agent(tmp_path)
    monkeypatch.setattr(api, "agent", selected)
    response = TestClient(api.app).post(
        "/v1/chat/completions",
        json={
            "model": api.MODEL_ID,
            "knowledge_id": "project-delivery",
            "messages": [{"role": "user", "content": "项目制度是什么？"}],
            "stream": True,
            "dry_run": True,
            "thinking": True,
        },
    )

    assert response.status_code == 200
    assert "reasoning_content" in response.text
    assert "data: [DONE]" in response.text


def test_api_emits_delta_files_only_from_query_sources(
    monkeypatch,
    tmp_path,
) -> None:
    content = b"%PDF-1.7\n%%EOF\n"
    digest = hashlib.sha256(content).hexdigest()
    documents_dir = tmp_path / "documents"
    documents_dir.mkdir()
    (documents_dir / "制度.pdf").write_bytes(content)
    source = SourceDocument(
        filename="制度.pdf",
        sha256=digest,
    )

    class SourceAgent:
        runtime = SimpleNamespace(
                settings=SimpleNamespace(
                    stream_heartbeat_seconds=0.01,
                    max_files_per_request=100,
                    max_download_files=10,
                    max_download_bytes=50 * 1024 * 1024,
                    max_download_image_bytes=10 * 1024 * 1024,
            ),
            manager=SimpleNamespace(
                active_documents_dir=lambda _knowledge_id: documents_dir
            ),
        )

        def invoke(self, **_kwargs):
            return AgentResult(
                intent=Intent.QUERY,
                content="回答\n\n来源：\n- 制度.pdf",
                knowledge_id="marketing",
                department="市场经营部",
                source_documents=[source],
            )

    monkeypatch.setattr(api, "agent", SourceAgent())
    client = TestClient(api.app)
    payload = {
        "model": api.MODEL_ID,
        "knowledge_id": "marketing",
        "messages": [{"role": "user", "content": "制度是什么？"}],
    }

    non_stream = client.post("/v1/chat/completions", json=payload).json()
    files = non_stream["choices"][0]["message"]["files"]
    assert len(files) == 1
    assert files[0]["filename"] == "制度.pdf"
    assert files[0]["file_type"] == "pdf"
    assert files[0]["sha256"] == digest
    assert base64.b64decode(files[0]["content_base64"]) == content

    stream = client.post(
        "/v1/chat/completions",
        json={**payload, "stream": True},
    )
    assert '"file"' in stream.text
    assert '"filename": "制度.pdf"' in stream.text
    assert '"content_base64"' in stream.text


def test_api_stream_emits_heartbeat_during_long_step(monkeypatch) -> None:
    class SlowAgent:
        runtime = SimpleNamespace(
            settings=SimpleNamespace(
                stream_heartbeat_seconds=0.01,
                max_files_per_request=100,
            )
        )

        def invoke(self, *, progress, **kwargs):
            progress(ProgressEvent("ocr", "正在 OCR：scan.pdf，第 1/2 页。"))
            time.sleep(0.035)
            return AgentResult(
                intent=Intent.SAVE,
                content="保存完成",
                knowledge_id="procurement-implementation",
                department="采购实施部",
            )

    monkeypatch.setattr(api, "agent", SlowAgent())
    response = TestClient(api.app).post(
        "/v1/chat/completions",
        json={
            "model": api.MODEL_ID,
            "knowledge_id": "procurement-implementation",
            "messages": [{"role": "user", "content": "请保存"}],
            "stream": True,
            "thinking": True,
        },
    )

    assert "正在 OCR：scan.pdf，第 1/2 页。" in response.text
    assert "仍在处理，当前阶段：ocr。" in response.text
    assert "保存完成" in response.text


def test_api_stream_reports_uncommitted_failure(monkeypatch) -> None:
    class FailingAgent:
        runtime = SimpleNamespace(
            settings=SimpleNamespace(
                stream_heartbeat_seconds=0.01,
                max_files_per_request=100,
            )
        )

        def invoke(self, *, progress, **kwargs):
            progress(ProgressEvent("embed", "正在生成向量。"))
            raise RuntimeError("internal details")

    monkeypatch.setattr(api, "agent", FailingAgent())
    response = TestClient(api.app).post(
        "/v1/chat/completions",
        json={
            "model": api.MODEL_ID,
            "knowledge_id": "procurement-implementation",
            "messages": [{"role": "user", "content": "请保存"}],
            "stream": True,
            "thinking": True,
        },
    )

    assert "本次处理失败，索引未提交。" in response.text
    assert "此前已发布的知识库快照保持不变" in response.text
    assert "internal details" not in response.text
    assert "data: [DONE]" in response.text


def test_traceback_redaction_removes_signed_query() -> None:
    value = (
        "failed https://minio.example/file.pdf?"
        "X-Amz-Signature=secret&X-Amz-Credential=value"
    )

    redacted = api._redact_traceback(value)

    assert "secret" not in redacted
    assert "Credential" not in redacted
    assert redacted.endswith("?[REDACTED]")


def test_signed_attachment_url_is_not_sent_as_intent_text() -> None:
    request = api.ChatCompletionRequest(
        model=api.MODEL_ID,
        knowledge_id="marketing",
        messages=[
            {
                "role": "user",
                "content": (
                    "请保存："
                    "https://minio.example/file.pdf?X-Amz-Signature=secret"
                ),
            }
        ],
    )

    text, sources = api._last_user_input(request)

    assert "X-Amz-Signature" not in text
    assert "[附件URL]" in text
    assert sources[0].url.endswith("X-Amz-Signature=secret")


def test_minio_archive_uses_department_bucket_and_content_address(
    monkeypatch,
) -> None:
    class MissingObject(Exception):
        code = "NoSuchKey"

    class FakeMinioClient:
        def __init__(self) -> None:
            self.buckets: set[str] = set()
            self.puts: list[tuple[str, str, bytes]] = []

        def bucket_exists(self, bucket: str) -> bool:
            return bucket in self.buckets

        def make_bucket(self, bucket: str) -> None:
            self.buckets.add(bucket)

        def stat_object(self, bucket: str, object_key: str):
            raise MissingObject

        def put_object(self, bucket, object_key, stream, *, length, **kwargs):
            self.puts.append((bucket, object_key, stream.read(length)))

    client = FakeMinioClient()
    monkeypatch.setattr("minio.Minio", lambda *args, **kwargs: client)
    store = MinioDepartmentObjectStore(
        DepartmentKnowledgeBaseSettings(
            minio_endpoint="minio:9000",
            minio_access_key="access",
            minio_secret_key="secret-key",
        )
    )
    from src.agents.department_knowledge_base.storage import PreparedDocument

    locations = store.archive(
        get_department("finance"),
        [
            PreparedDocument(
                filename="policy.pdf",
                data=b"pdf",
                sha256="a" * 64,
            )
        ],
    )

    assert locations[0].bucket == "department-kb-finance"
    assert locations[0].object_key == (
        f"sha256/aa/{'a' * 64}/policy.pdf"
    )
    assert client.puts[0][2] == b"pdf"
