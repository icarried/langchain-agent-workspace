from __future__ import annotations

import base64
from pathlib import Path

import pytest
from docx import Document
from fastapi.testclient import TestClient
from fastmcp import Client
from typer.testing import CliRunner

from src.agents.contract_review.api import app
from src.agents.contract_review.chunking import chunk_elements
from src.agents.contract_review.cli import app as cli_app
from src.agents.contract_review.contract_loader import load_contract_elements
from src.agents.contract_review.graph import build_graph
from src.agents.contract_review.mcp_server import mcp
from src.agents.contract_review.service import review_contract

runner = CliRunner()


def test_txt_loader_extracts_contract_lines(tmp_path: Path) -> None:
    contract = tmp_path / "contract.txt"
    contract.write_text("第一条 服务内容\n乙方提供系统开发服务\n第五条 违约责任\n", encoding="utf-8")

    elements = load_contract_elements(contract)

    assert [element.kind for element in elements] == ["paragraph", "paragraph", "paragraph"]
    assert elements[0].text == "第一条 服务内容"


def test_docx_loader_extracts_paragraphs_and_tables(tmp_path: Path) -> None:
    contract = _make_docx_contract(tmp_path)

    elements = load_contract_elements(contract)

    assert elements[0].text == "技术服务合同"
    assert any(element.kind == "table" and "付款 | 上线后支付尾款" in element.text for element in elements)


def test_chunk_elements_splits_on_contract_sections(tmp_path: Path) -> None:
    elements = load_contract_elements(_make_docx_contract(tmp_path))

    chunks = chunk_elements(elements, max_chars=500)

    assert chunks
    assert chunks[0].chunk_id == "chunk-001"
    assert any("第一条" in chunk.title for chunk in chunks)


def test_graph_dry_run_creates_contract_report(tmp_path: Path) -> None:
    contract = tmp_path / "contract.txt"
    contract.write_text(
        "甲方：A公司\n乙方：B公司\n第一条 服务内容\n乙方提供数据分析平台。\n第五条 违约责任\n任何一方违约承担责任。\n",
        encoding="utf-8",
    )
    output = tmp_path / "report.md"
    graph = build_graph()

    result = graph.invoke(
        {
            "contract_path": str(contract),
            "client_role": "甲方",
            "contract_type": "技术服务合同",
            "transaction_background": "采购数据分析平台",
            "review_guide_path": str(Path("src/agents/contract_review/review_guide/合同审查六维规则.md")),
            "output_path": str(output),
            "dry_run": True,
            "provider": "deepseek",
        }
    )

    assert "合同审查 dry-run 报告" in result["final_report"]
    assert "六个维度" in result["final_report"]
    assert "甲方" in result["final_report"]
    assert output.exists()


def test_service_dry_run_with_builtin_example() -> None:
    root = Path(__file__).resolve().parents[2]
    contract = root / "src" / "agents" / "contract_review" / "examples" / "示例服务合同.md"

    result = review_contract(
        contract,
        client_role="甲方",
        contract_type="技术服务合同",
        transaction_background="采购设备运行数据分析平台",
        dry_run=True,
    )

    assert result["dry_run"] is True
    assert result["chunk_count"] >= 1
    assert "FastGPT 合同审查大师" in result["report"]


def test_api_review_dry_run(tmp_path: Path) -> None:
    contract = tmp_path / "contract.txt"
    contract.write_text("第一条 服务内容\n乙方提供开发服务\n", encoding="utf-8")
    client = TestClient(app)

    health = client.get("/health")
    assert health.status_code == 200
    assert health.json()["agent"] == "contract-review"

    response = client.post(
        "/review",
        json={
            "contract_path": str(contract),
            "client_role": "乙方",
            "contract_type": "服务合同",
            "transaction_background": "乙方承接平台开发",
            "dry_run": True,
        },
    )

    assert response.status_code == 200
    data = response.json()
    assert data["dry_run"] is True
    assert data["chunk_count"] >= 1
    assert "合同审查 dry-run 报告" in data["report"]


def test_cli_review_subcommand_accepts_builtin_example() -> None:
    root = Path(__file__).resolve().parents[2]
    contract = root / "src" / "agents" / "contract_review" / "examples" / "示例服务合同.md"

    result = runner.invoke(
        cli_app,
        [
            "review",
            str(contract),
            "--client-role",
            "甲方",
            "--contract-type",
            "技术服务合同",
            "--dry-run",
        ],
    )

    assert result.exit_code == 0
    assert "合同审查 dry-run 报告" in result.output


@pytest.mark.asyncio
async def test_mcp_review_tool_accepts_base64_contract_upload() -> None:
    content = "第一条 服务内容\n乙方提供系统开发服务\n".encode("utf-8")

    async with Client(mcp) as client:
        tools = await client.list_tools()
        assert any(tool.name == "review_contract" for tool in tools)

        result = await client.call_tool(
            "review_contract",
            {
                "contract_base64": base64.b64encode(content).decode("ascii"),
                "contract_filename": "contract.txt",
                "client_role": "甲方",
                "contract_type": "技术服务合同",
                "dry_run": True,
            },
        )

    assert result.data["dry_run"] is True
    assert result.data["chunk_count"] >= 1
    assert result.data["filename"] == "contract.txt"
    assert "contract_path" not in result.data


def _make_docx_contract(tmp_path: Path) -> Path:
    document = Document()
    document.add_heading("技术服务合同", level=1)
    document.add_paragraph("第一条 服务内容")
    document.add_paragraph("乙方为甲方提供设备数据分析平台开发服务。")
    table = document.add_table(rows=2, cols=2)
    table.cell(0, 0).text = "事项"
    table.cell(0, 1).text = "约定"
    table.cell(1, 0).text = "付款"
    table.cell(1, 1).text = "上线后支付尾款"
    path = tmp_path / "contract.docx"
    document.save(path)
    return path
