from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import Cm

from src.agents.official_document_formatting.formatter import format_docx
from src.agents.official_document_formatting.graph import build_graph
from src.agents.official_document_formatting.service import format_official_document
import src.agents.official_document_formatting.service as formatting_service


def _make_company_sample(path: Path) -> Path:
    document = Document()
    section = document.sections[0]
    section.top_margin = Cm(1)
    section.bottom_margin = Cm(1)
    section.left_margin = Cm(1)
    section.right_margin = Cm(1)
    document.add_paragraph("关于开展测试工作的请示")
    document.add_paragraph("公司领导：")
    document.add_paragraph("现将有关情况报告如下。")
    document.add_paragraph("一、主要事项")
    document.add_paragraph("（一）具体安排")
    document.add_paragraph("请各部门按要求完成材料报送。")
    document.add_paragraph("请各部门按要求完成材料报送。")
    document.add_paragraph("附件：")
    document.add_paragraph("1. 测试材料清单")
    document.add_paragraph("技术支撑部")
    document.add_paragraph("2026年7月29日")
    table = document.add_table(rows=2, cols=4)
    table.cell(0, 0).text = "序号"
    table.cell(0, 1).text = "货物名称"
    table.cell(0, 2).text = "数量"
    table.cell(0, 3).text = "预算总价（元）"
    table.cell(1, 0).text = "1"
    table.cell(1, 1).text = "测试设备"
    table.cell(1, 2).text = "2"
    table.cell(1, 3).text = "12,000.00"
    document.save(path)
    return path


def _east_asia_font(paragraph) -> str | None:
    run = next(run for run in paragraph.runs if run.text.strip())
    fonts = run._r.get_or_add_rPr().rFonts
    return fonts.get(qn("w:eastAsia"))


def test_formatter_applies_v1_company_rules_without_changing_content(
    tmp_path: Path,
) -> None:
    source = _make_company_sample(tmp_path / "input.docx")
    output = tmp_path / "output.docx"

    format_docx(source, output)

    formatted = Document(output)
    section = formatted.sections[0]
    assert round(section.top_margin.cm, 1) == 3.7
    assert round(section.bottom_margin.cm, 1) == 3.5
    assert round(section.left_margin.cm, 1) == 2.8
    assert round(section.right_margin.cm, 1) == 2.6
    assert round(section.page_width.cm, 1) == 21.0
    assert round(section.page_height.cm, 1) == 29.7
    texts = [paragraph.text for paragraph in formatted.paragraphs if paragraph.text.strip()]
    assert texts == [
        "关于开展测试工作的请示",
        "公司领导：",
        "现将有关情况报告如下。",
        "一、主要事项",
        "（一）具体安排",
        "请各部门按要求完成材料报送。",
        "请各部门按要求完成材料报送。",
        "附件：",
        "1. 测试材料清单",
        "技术支撑部",
        "2026年7月29日",
    ]
    non_empty = [paragraph for paragraph in formatted.paragraphs if paragraph.text.strip()]
    assert non_empty[0].alignment == WD_ALIGN_PARAGRAPH.CENTER
    assert _east_asia_font(non_empty[0]) == "方正小标宋简体"
    assert _east_asia_font(non_empty[1]) == "仿宋_GB2312"
    assert _east_asia_font(non_empty[3]) == "黑体"
    assert _east_asia_font(non_empty[4]) == "楷体_GB2312"
    assert _east_asia_font(non_empty[5]) == "仿宋_GB2312"
    assert non_empty[1].alignment == WD_ALIGN_PARAGRAPH.LEFT
    assert non_empty[7].alignment == WD_ALIGN_PARAGRAPH.LEFT
    assert non_empty[8].alignment == WD_ALIGN_PARAGRAPH.LEFT
    assert non_empty[9].alignment == WD_ALIGN_PARAGRAPH.RIGHT
    assert non_empty[10].alignment == WD_ALIGN_PARAGRAPH.RIGHT
    assert non_empty[1]._p.get_or_add_pPr().ind.get(qn("w:firstLine")) is None
    assert non_empty[3]._p.get_or_add_pPr().ind.get(qn("w:firstLine")) == "640"
    assert (
        non_empty[3]._p.get_or_add_pPr().ind.get(qn("w:firstLineChars"))
        == "200"
    )
    assert non_empty[4]._p.get_or_add_pPr().ind.get(qn("w:firstLine")) == "640"
    assert non_empty[8]._p.get_or_add_pPr().ind.get(qn("w:leftChars")) == "200"
    body_properties = non_empty[5]._p.get_or_add_pPr()
    assert body_properties.ind.get(qn("w:firstLine")) == "640"
    assert body_properties.ind.get(qn("w:firstLineChars")) == "200"
    assert body_properties.spacing.get(qn("w:line")) == "560"
    assert body_properties.spacing.get(qn("w:lineRule")) == "exact"
    borders = formatted.tables[0]._tbl.tblPr.find(qn("w:tblBorders"))
    assert borders.find(qn("w:top")).get(qn("w:val")) == "single"
    assert borders.find(qn("w:top")).get(qn("w:sz")) == "12"
    assert borders.find(qn("w:bottom")).get(qn("w:sz")) == "12"
    assert borders.find(qn("w:left")).get(qn("w:val")) == "none"
    assert borders.find(qn("w:right")).get(qn("w:val")) == "none"
    assert borders.find(qn("w:insideH")).get(qn("w:val")) == "none"
    assert borders.find(qn("w:insideV")).get(qn("w:val")) == "none"
    header_row_properties = formatted.tables[0].rows[0]._tr.get_or_add_trPr()
    assert header_row_properties.find(qn("w:tblHeader")) is not None
    header_cell_borders = formatted.tables[0].cell(0, 0)._tc.tcPr.find(
        qn("w:tcBorders")
    )
    assert header_cell_borders.find(qn("w:bottom")).get(qn("w:sz")) == "6"
    assert all(
        row._tr.get_or_add_trPr().find(qn("w:cantSplit")) is not None
        for row in formatted.tables[0].rows
    )
    body_row = formatted.tables[0].rows[1]
    assert body_row.cells[0].paragraphs[0].alignment == WD_ALIGN_PARAGRAPH.CENTER
    assert body_row.cells[1].paragraphs[0].alignment == WD_ALIGN_PARAGRAPH.LEFT
    assert body_row.cells[2].paragraphs[0].alignment == WD_ALIGN_PARAGRAPH.CENTER
    assert body_row.cells[3].paragraphs[0].alignment == WD_ALIGN_PARAGRAPH.RIGHT


def test_graph_dry_run_validates_without_writing_output(tmp_path: Path) -> None:
    source = _make_company_sample(tmp_path / "input.docx")
    output = tmp_path / "output.docx"

    result = build_graph().invoke(
        {
            "source_path": str(source),
            "output_path": str(output),
            "original_filename": source.name,
            "dry_run": True,
        }
    )

    assert result["result"].dry_run is True
    assert result["result"].content == b""
    assert not output.exists()


def test_service_returns_formatted_docx_bytes_and_keeps_source(tmp_path: Path) -> None:
    source = _make_company_sample(tmp_path / "input.docx")
    original = source.read_bytes()

    result = format_official_document(source)

    assert source.read_bytes() == original
    assert result["filename"] == "input-公文格式化.docx"
    assert result["mime_type"].endswith("wordprocessingml.document")
    assert result["size"] == len(result["content"])
    assert len(result["sha256"]) == 64
    assert result["content"].startswith(b"PK")
    assert result["output_path"] == ""
    assert "font_status" not in result
    assert "available_fonts" not in result
    assert "missing_fonts" not in result
    assert result["findings"]
    assert all(
        {"severity", "rule_id", "paragraph_index", "message", "verified"}
        <= finding.keys()
        for finding in result["findings"]
    )
    assert any(not finding["verified"] for finding in result["findings"])
    assert "未验证" in result["report"]


def test_service_rejects_non_docx_input(tmp_path: Path) -> None:
    source = tmp_path / "input.txt"
    source.write_text("not a docx", encoding="utf-8")

    try:
        format_official_document(source)
    except ValueError as exc:
        assert "DOCX" in str(exc)
    else:
        raise AssertionError("non-DOCX input should be rejected")


def test_service_converts_legacy_doc_then_returns_docx(
    tmp_path: Path,
    monkeypatch,
) -> None:
    converted = _make_company_sample(tmp_path / "converted.docx").read_bytes()
    source = tmp_path / "input.doc"
    source.write_bytes(b"legacy-doc-placeholder")
    monkeypatch.setattr(
        formatting_service,
        "convert_doc_to_docx",
        lambda data, *, source: converted,
    )

    result = format_official_document(source)

    assert result["filename"] == "input-公文格式化.docx"
    assert result["content"].startswith(b"PK")
