from __future__ import annotations

import base64
from pathlib import Path

import pytest
from docx import Document
from docx.shared import Mm
from fastapi.testclient import TestClient
from fastmcp import Client
from typer.testing import CliRunner

from src.agents.official_document_review.api import app
from src.agents.official_document_review.cli import app as cli_app
from src.agents.official_document_review.document_loader import load_document_elements
from src.agents.official_document_review.format_checks import inspect_official_document
from src.agents.official_document_review.graph import build_graph
from src.agents.official_document_review.mcp_server import mcp
from src.agents.official_document_review.service import review_official_document

runner = CliRunner()


def test_txt_loader_extracts_document_lines(tmp_path: Path) -> None:
    document = tmp_path / "notice.txt"
    document.write_text("关于开展测试工作的通知\n各部门：\n请按时报送材料。\n", encoding="utf-8")

    elements = load_document_elements(document)

    assert [element.kind for element in elements] == ["paragraph", "paragraph", "paragraph"]
    assert elements[0].text == "关于开展测试工作的通知"


def test_docx_inspector_flags_bad_margins(tmp_path: Path) -> None:
    document = _make_docx_notice(tmp_path, bad_margins=True)
    elements = load_document_elements(document)

    findings = inspect_official_document(document, elements)

    assert any(finding.check_id.startswith("page-") for finding in findings)
    assert any("成文日期" in finding.category for finding in findings)


def test_graph_dry_run_creates_report(tmp_path: Path) -> None:
    document = tmp_path / "notice.txt"
    document.write_text("关于开展测试工作的通知\n各部门：\n请按时报送材料。\n", encoding="utf-8")
    output = tmp_path / "report.md"
    graph = build_graph()

    result = graph.invoke(
        {
            "document_path": str(document),
            "document_type": "通知",
            "review_guide_path": str(Path("src/agents/official_document_review/review_guide/党政机关公文格式检查要点.md")),
            "output_path": str(output),
            "dry_run": True,
            "provider": "deepseek",
        }
    )

    assert "公文格式检查 dry-run 报告" in result["final_report"]
    assert "FastGPT 公文优化" in result["final_report"]
    assert output.exists()


def test_service_dry_run_with_builtin_example() -> None:
    root = Path(__file__).resolve().parents[2]
    document = root / "src" / "agents" / "official_document_review" / "examples" / "示例通知.md"

    result = review_official_document(document, document_type="通知", dry_run=True)

    assert result["dry_run"] is True
    assert result["finding_count"] >= 1
    assert "公文格式检查 dry-run 报告" in result["report"]


def test_api_review_dry_run(tmp_path: Path) -> None:
    document = tmp_path / "notice.txt"
    document.write_text("关于开展测试工作的通知\n各部门：\n请落实。\n", encoding="utf-8")
    client = TestClient(app)

    health = client.get("/health")
    assert health.status_code == 200
    assert health.json()["agent"] == "official-document-review"

    response = client.post(
        "/review",
        json={"document_path": str(document), "document_type": "通知", "dry_run": True},
    )

    assert response.status_code == 200
    data = response.json()
    assert data["dry_run"] is True
    assert data["finding_count"] >= 1
    assert "公文格式检查 dry-run 报告" in data["report"]


def test_cli_review_subcommand_accepts_builtin_example() -> None:
    root = Path(__file__).resolve().parents[2]
    document = root / "src" / "agents" / "official_document_review" / "examples" / "示例通知.md"

    result = runner.invoke(cli_app, ["review", str(document), "--document-type", "通知", "--dry-run"])

    assert result.exit_code == 0
    assert "公文格式检查 dry-run 报告" in result.output


@pytest.mark.asyncio
async def test_mcp_review_tool_accepts_base64_document_upload() -> None:
    content = "关于开展测试工作的通知\n各部门：\n请落实。\n".encode("utf-8")

    async with Client(mcp) as client:
        tools = await client.list_tools()
        assert any(tool.name == "review_official_document" for tool in tools)

        result = await client.call_tool(
            "review_official_document",
            {
                "document_base64": base64.b64encode(content).decode("ascii"),
                "document_filename": "notice.txt",
                "document_type": "通知",
                "dry_run": True,
            },
        )

    assert result.data["dry_run"] is True
    assert result.data["finding_count"] >= 1
    assert result.data["filename"] == "notice.txt"
    assert "document_path" not in result.data


def _make_docx_notice(tmp_path: Path, *, bad_margins: bool) -> Path:
    document = Document()
    section = document.sections[0]
    if bad_margins:
        section.top_margin = Mm(20)
        section.bottom_margin = Mm(20)
        section.left_margin = Mm(20)
        section.right_margin = Mm(20)
    document.add_heading("关于开展测试工作的通知", level=1)
    document.add_paragraph("各部门：")
    document.add_paragraph("请按时报送材料。")
    path = tmp_path / "notice.docx"
    document.save(path)
    return path
