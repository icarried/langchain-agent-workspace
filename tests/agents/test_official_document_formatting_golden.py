from __future__ import annotations

import hashlib
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn

from src.agents.official_document_formatting.compliance import evaluate_compliance
from src.agents.official_document_formatting.formatter import (
    document_semantic_snapshot,
    format_docx,
)
from src.agents.official_document_formatting.roles import (
    ParagraphRole,
    classify_paragraphs,
)
from src.agents.official_document_formatting.standards import FS, FZXBS, HEI

WORKSPACE = Path(__file__).resolve().parents[2]
FIXTURES = WORKSPACE / "临时文件" / "公文格式化配置"
STANDARD = FIXTURES / "公文格式规范.docx"
PURCHASE_REQUEST = (
    FIXTURES
    / "关于采购多场景机器人及具身智能技术研发及能力建设项目配套设备的请示-原版.docx"
)


def _east_asia_font(paragraph) -> str | None:
    run = next(run for run in paragraph.runs if run.text.strip())
    return run._r.get_or_add_rPr().rFonts.get(qn("w:eastAsia"))


def test_normative_standard_text_snapshot_is_stable() -> None:
    document = Document(STANDARD)
    texts = [paragraph.text.strip() for paragraph in document.paragraphs if paragraph.text.strip()]
    digest = hashlib.sha256("\n".join(texts).encode("utf-8")).hexdigest()

    assert len(document.paragraphs) == 34
    assert len(texts) == 30
    assert len(document.tables) == 0
    assert digest == "8b73a08404fc8444b20e96189e696dc5f4806242882b60f319bbd951a71bc30a"


def test_purchase_request_matches_structural_golden_rules(tmp_path: Path) -> None:
    output = tmp_path / "purchase-request-formatted.docx"
    source = Document(PURCHASE_REQUEST)

    format_docx(PURCHASE_REQUEST, output)

    formatted = Document(output)
    assert document_semantic_snapshot(formatted) == document_semantic_snapshot(source)

    section = formatted.sections[0]
    assert (
        round(section.top_margin.cm, 1),
        round(section.bottom_margin.cm, 1),
        round(section.left_margin.cm, 1),
        round(section.right_margin.cm, 1),
    ) == (3.7, 3.5, 2.8, 2.6)
    assert round(section.footer_distance.cm, 1) == 2.5
    assert "PAGE" in section.footer._element.xml
    assert "PAGE" in section.even_page_footer._element.xml
    grid = section._sectPr.find(qn("w:docGrid"))
    assert grid.get(qn("w:linePitch")) == "560"

    roles = classify_paragraphs([paragraph.text for paragraph in formatted.paragraphs])
    assert roles.count(ParagraphRole.HEADING_1) == 6
    title = formatted.paragraphs[roles.index(ParagraphRole.TITLE)]
    recipient = formatted.paragraphs[roles.index(ParagraphRole.MAIN_RECIPIENT)]
    heading = formatted.paragraphs[roles.index(ParagraphRole.HEADING_1)]
    attachment = formatted.paragraphs[roles.index(ParagraphRole.ATTACHMENT_NOTE)]
    attachment_item = formatted.paragraphs[roles.index(ParagraphRole.ATTACHMENT_ITEM)]
    signature = formatted.paragraphs[roles.index(ParagraphRole.SIGNATURE)]
    date = formatted.paragraphs[roles.index(ParagraphRole.DATE)]

    assert title.alignment == WD_ALIGN_PARAGRAPH.CENTER
    assert _east_asia_font(title) == FZXBS
    assert recipient.alignment == WD_ALIGN_PARAGRAPH.LEFT
    assert _east_asia_font(recipient) == FS
    assert _east_asia_font(heading) == HEI
    assert heading._p.pPr.ind.get(qn("w:firstLineChars")) == "200"
    assert attachment.text == "附件："
    assert attachment._p.pPr.ind.get(qn("w:firstLineChars")) == "200"
    assert not attachment_item.text.startswith((" ", "\t", "\u3000"))
    assert attachment_item._p.pPr.ind.get(qn("w:leftChars")) == "200"
    assert signature._p.pPr.ind.get(qn("w:rightChars")) == "400"
    assert date._p.pPr.ind.get(qn("w:rightChars")) == "400"

    table = formatted.tables[0]
    body_row = table.rows[1]
    assert body_row.cells[0].paragraphs[0].alignment == WD_ALIGN_PARAGRAPH.CENTER
    assert body_row.cells[1].paragraphs[0].alignment == WD_ALIGN_PARAGRAPH.LEFT
    assert body_row.cells[2].paragraphs[0].alignment == WD_ALIGN_PARAGRAPH.LEFT
    assert body_row.cells[3].paragraphs[0].alignment == WD_ALIGN_PARAGRAPH.CENTER
    assert body_row.cells[4].paragraphs[0].alignment == WD_ALIGN_PARAGRAPH.CENTER
    assert body_row.cells[-1].paragraphs[0].alignment == WD_ALIGN_PARAGRAPH.RIGHT

    findings = evaluate_compliance(formatted, formatted=True)
    assert not [finding for finding in findings if finding.verified]
    assert any(finding.rule_id == "visual.pagination" for finding in findings)
    assert any(finding.rule_id == "visual.seal-position" for finding in findings)
